#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
from datetime import datetime
from pathlib import Path


def normalize_heading_text_sau_fbe(value: str) -> str:
    value = (value or "").strip().lower()
    table = str.maketrans({
        "İ": "i", "I": "i", "ı": "i",
        "Ş": "s", "ş": "s",
        "Ğ": "g", "ğ": "g",
        "Ü": "u", "ü": "u",
        "Ö": "o", "ö": "o",
        "Ç": "c", "ç": "c",
    })
    return value.translate(table)


def find_heading_position_sau_fbe(all_texts, heading: str):
    aliases = {
        "Özet": ["özet"],
        "Abstract": ["abstract", "summary"],
        "İçindekiler": ["içindekiler", "icindekiler"],
        "Tablo Listesi": ["tablo listesi"],
        "Şekil Listesi": ["şekil listesi", "sekil listesi"],
        "Kaynakça": ["kaynakça", "kaynaklar"],
    }

    normalized_texts = [normalize_heading_text_sau_fbe(t) for t in all_texts]
    search_terms = aliases.get(heading, [heading])
    search_terms = [normalize_heading_text_sau_fbe(x) for x in search_terms]

    for idx, text in enumerate(normalized_texts):
        if any(term in text for term in search_terms):
            return idx
    return None


def recompute_expected_headings_sau_fbe(all_texts, expected_headings):
    positions = {}
    missing = []

    for heading in expected_headings:
        pos = find_heading_position_sau_fbe(all_texts, heading)
        positions[heading] = pos
        if pos is None:
            missing.append(heading)

    body_headings = [
        "Bölüm 1. Giriş",
        "Bölüm 2. Kavramsal ve Kuramsal Arka Plan",
        "Bölüm 3. Literatür Taraması ve İlgili Çalışmalar",
        "Bölüm 4. Yöntem ve Çalışma Zamanı Doğrulama",
        "Bölüm 5. Tartışma, Sınırlılıklar ve Gelecek Çalışmalar",
        "Bölüm 6. Sonuç ve Öneriler",
    ]

    body_positions = [positions.get(h) for h in body_headings if positions.get(h) is not None]
    heading_order_ok = body_positions == sorted(body_positions)

    return positions, missing, heading_order_ok



def normalize_tr_text_for_heading_check(value: str) -> str:
    value = (value or "").strip().lower()
    table = str.maketrans({
        "İ": "i",
        "I": "i",
        "ı": "i",
        "Ş": "s",
        "ş": "s",
        "Ğ": "g",
        "ğ": "g",
        "Ü": "u",
        "ü": "u",
        "Ö": "o",
        "ö": "o",
        "Ç": "c",
        "ç": "c",
    })
    return value.translate(table)


def find_heading_position_case_insensitive(all_texts, heading: str):
    aliases = {
        "Abstract": ["abstract", "summary"],
        "İçindekiler": ["içindekiler", "icindekiler"],
        "Tablo Listesi": ["tablo listesi"],
        "Şekil Listesi": ["şekil listesi", "sekil listesi"],
        "Kaynakça": ["kaynakça", "kaynaklar"],
    }
    normalized_texts = [normalize_tr_text_for_heading_check(t) for t in all_texts]
    search_terms = aliases.get(heading, [heading])
    search_terms = [normalize_tr_text_for_heading_check(x) for x in search_terms]
    for idx, text in enumerate(normalized_texts):
        if any(term in text for term in search_terms):
            return idx
    return None


def recompute_heading_status_case_insensitive(all_texts, expected_headings):
    positions = {}
    missing = []
    for h in expected_headings:
        pos = find_heading_position_case_insensitive(all_texts, h)
        if pos is None:
            missing.append(h)
            positions[h] = None
        else:
            positions[h] = pos

    # Body chapter order is the critical order. Frontmatter may be generated separately
    # and list placeholders can appear before numbered chapters.
    body = [
        "Bölüm 1. Giriş",
        "Bölüm 2. Kavramsal ve Kuramsal Arka Plan",
        "Bölüm 3. Literatür Taraması ve İlgili Çalışmalar",
        "Bölüm 4. Yöntem ve Çalışma Zamanı Doğrulama",
        "Bölüm 5. Tartışma, Sınırlılıklar ve Gelecek Çalışmalar",
        "Bölüm 6. Sonuç ve Öneriler",
    ]
    body_positions = [positions.get(h) for h in body if positions.get(h) is not None]
    body_order_ok = body_positions == sorted(body_positions)

    return positions, missing, body_order_ok



def append_sau_fbe_precheck_section(md, doc, all_texts):
    """Append preliminary SAÜ FBE thesis guide/template compliance checks."""
    md.append("")
    md.append("## SAÜ FBE Format Uyum Ön Kontrolü")
    md.append("")
    md.append("Bu bölüm, Sakarya Üniversitesi Fen Bilimleri Enstitüsü Lisansüstü Tez Yazım Kılavuzu ve Tez Hazırlama Şablonu esas alınarak yapılan ön format kontrolüdür.")
    md.append("")
    md.append("| Kontrol | Durum | Not |")
    md.append("|---|---|---|")

    fbe_guide_pdf = Path("docs/thesis_template_sources/fbe_lisansustu_tez_yazim_kilavuzu.pdf")
    fbe_guide_txt = Path("docs/thesis_template_sources/fbe_lisansustu_tez_yazim_kilavuzu.txt")
    fbe_template = Path("docs/thesis_template_sources/tez_hazirlama_sablonu.docx")
    fbe_rules = Path("docs/tez_format_kontrol_kriterleri.md")

    md.append(f"| FBE tez yazım kılavuzu PDF | {'ok' if fbe_guide_pdf.exists() else 'missing'} | `{fbe_guide_pdf}` |")
    md.append(f"| FBE tez yazım kılavuzu TXT | {'ok' if fbe_guide_txt.exists() else 'missing'} | `{fbe_guide_txt}` |")
    md.append(f"| Tez hazırlama şablonu DOCX | {'ok' if fbe_template.exists() else 'missing'} | `{fbe_template}` |")
    md.append(f"| Format kontrol kriterleri | {'ok' if fbe_rules.exists() else 'missing'} | `{fbe_rules}` |")

    all_texts_lower = [t.lower() for t in all_texts]

    frontmatter_expected = [
        "Özet",
        "Abstract",
        "İçindekiler",
        "Tablo Listesi",
        "Şekil Listesi",
        "Kaynakça",
    ]

    aliases = {
        "Abstract": ["abstract", "summary"],
        "İçindekiler": ["içindekiler", "icindekiler"],
        "Tablo Listesi": ["tablo listesi"],
        "Şekil Listesi": ["şekil listesi", "sekil listesi"],
    }

    for item in frontmatter_expected:
        search_terms = aliases.get(item, [item.lower()])
        status = "ok" if any(any(term in t for term in search_terms) for t in all_texts_lower) else "check"
        md.append(f"| Ön bölüm/metin: {item} | {status} | Ana DOCX içinde arandı |")

    table_status = "ok" if len(doc.tables) >= 1 else "check"
    shape_status = "ok" if len(doc.inline_shapes) >= 1 else "check"
    md.append(f"| Word tablo nesnesi | {table_status} | Tablo sayısı: {len(doc.tables)} |")
    md.append(f"| Word şekil nesnesi | {shape_status} | Inline shape sayısı: {len(doc.inline_shapes)} |")

    try:
        section = doc.sections[0]
        margin_note = (
            f"top={section.top_margin.cm:.2f} cm, "
            f"bottom={section.bottom_margin.cm:.2f} cm, "
            f"left={section.left_margin.cm:.2f} cm, "
            f"right={section.right_margin.cm:.2f} cm"
        )
        md.append(f"| Sayfa kenar boşlukları okunabilirliği | ok | {margin_note} |")
    except Exception as exc:
        md.append(f"| Sayfa kenar boşlukları okunabilirliği | check | {exc} |")

    try:
        normal = doc.styles["Normal"]
        font_name = normal.font.name or "not-set"
        font_size = normal.font.size.pt if normal.font.size else "not-set"
        md.append(f"| Normal stil yazı tipi/punto | ok | font={font_name}, size={font_size} |")
    except Exception as exc:
        md.append(f"| Normal stil yazı tipi/punto | check | {exc} |")

    md.append("")
    md.append("Not: Bu bölüm şu anda ön format kontrolüdür. Bir sonraki aşamada kenar boşluğu, yazı tipi, satır aralığı, tablo/şekil başlığı konumu ve kaynakça biçimi kılavuzdaki kesin değerlere göre pass/fail kontrolüne dönüştürülecektir.")


from docx import Document


FRONTMATTER_HEADINGS = [
    "Özet",
    "Abstract",
    "İçindekiler",
    "Tablo Listesi",
    "Şekil Listesi",
]

CHAPTER_HEADINGS = [
    "Bölüm 1. Giriş",
    "Bölüm 2. Kavramsal ve Kuramsal Arka Plan",
    "Bölüm 3. Literatür Taraması ve İlgili Çalışmalar",
    "Bölüm 4. Yöntem ve Çalışma Zamanı Doğrulama",
    "Bölüm 5. Tartışma, Sınırlılıklar ve Gelecek Çalışmalar",
    "Bölüm 6. Sonuç ve Öneriler",
]

EXPECTED_HEADINGS = FRONTMATTER_HEADINGS + CHAPTER_HEADINGS

EXPECTED_FILES = [
    "docs/bolum_1_giris_tr.md",
    "docs/bolum_1_giris_tr.docx",
    "docs/bolum_2_kavramsal_kuramsal_arka_plan_tr.md",
    "docs/bolum_2_kavramsal_kuramsal_arka_plan_tr.docx",
    "docs/bolum_3_literatur_taramasi_ve_ilgili_calismalar_tr.md",
    "docs/bolum_3_literatur_taramasi_ve_ilgili_calismalar_tr.docx",
    "docs/bolum_4_yontem_ve_runtime_dogrulama_tr.md",
    "docs/bolum_4_yontem_ve_runtime_dogrulama_tr_with_artifacts.docx",
    "docs/bolum_5_tartisma_sinirliliklar_gelecek_calismalar_tr.md",
    "docs/bolum_5_tartisma_sinirliliklar_gelecek_calismalar_tr.docx",
    "docs/bolum_6_sonuc_ve_oneriler_tr.md",
    "docs/bolum_6_sonuc_ve_oneriler_tr.docx",
    "docs/tez_ana_taslak_tr.docx",
    "docs/tez_ana_taslak_tr_frontmatter.docx",
    "docs/tez_ana_taslak_tr_frontmatter_ozetli.docx",
    "docs/tez_dosya_envanteri.md",
    "docs/tez_bolum_paketi_readme.md",
    "docs/roadmap.md",
]

EXPECTED_ARTIFACTS = [
    "experiments/results/thesis_artifacts/run_05_port_aware_runtime_validation/tables/table_canonical_runtime_validation_summary.csv",
    "experiments/results/thesis_artifacts/run_05_port_aware_runtime_validation/tables/table_controller_action_distribution.csv",
    "experiments/results/thesis_artifacts/run_05_port_aware_runtime_validation/tables/table_final_top20_prediction_distribution.csv",
    "experiments/results/thesis_artifacts/run_05_port_aware_runtime_validation/tables/table_protocol_aware_final_policy_distribution.csv",
    "experiments/results/thesis_artifacts/run_05_port_aware_runtime_validation/tables/table_enforcement_action_summary.csv",
    "experiments/results/thesis_artifacts/run_05_port_aware_runtime_validation/tables/table_flow_level_model_controller_comparison.csv",
    "experiments/results/thesis_artifacts/run_05_port_aware_runtime_validation/tables/table_protocol_aware_attack_probability_summary.csv",
    "experiments/results/thesis_artifacts/run_05_port_aware_runtime_validation/figures/fig_controller_action_distribution.png",
    "experiments/results/thesis_artifacts/run_05_port_aware_runtime_validation/figures/fig_final_top20_prediction_distribution.png",
    "experiments/results/thesis_artifacts/run_05_port_aware_runtime_validation/figures/fig_protocol_aware_final_policy_distribution.png",
    "experiments/results/thesis_artifacts/run_05_port_aware_runtime_validation/figures/fig_enforcement_action_summary.png",
]


def file_size(path: Path) -> str:
    if not path.exists():
        return "-"
    n = path.stat().st_size
    if n < 1024:
        return f"{n} B"
    if n < 1024**2:
        return f"{n/1024:.1f} KB"
    return f"{n/1024**2:.1f} MB"


def normalize_heading_text(value: str) -> str:
    value = (value or "").strip().lower()
    table = str.maketrans({
        "İ": "i", "I": "i", "ı": "i",
        "Ş": "s", "ş": "s",
        "Ğ": "g", "ğ": "g",
        "Ü": "u", "ü": "u",
        "Ö": "o", "ö": "o",
        "Ç": "c", "ç": "c",
    })
    return value.translate(table)


def heading_aliases(heading: str) -> list[str]:
    aliases = {
        "Özet": ["özet"],
        "Abstract": ["abstract", "summary"],
        "İçindekiler": ["içindekiler", "icindekiler"],
        "Tablo Listesi": ["tablo listesi"],
        "Şekil Listesi": ["şekil listesi", "sekil listesi"],
        "Kaynakça": ["kaynakça", "kaynaklar"],
    }
    return aliases.get(heading, [heading])


def heading_matches(text: str, heading: str) -> bool:
    norm_text = normalize_heading_text(text)
    terms = [normalize_heading_text(x) for x in heading_aliases(heading)]
    return any(term in norm_text for term in terms)


def find_heading_positions(paragraphs: list[str], headings: list[str]) -> dict:
    result = {}
    for heading in headings:
        pos = None
        for i, text in enumerate(paragraphs):
            if heading_matches(text, heading):
                pos = i
                break
        result[heading] = pos
    return result


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--docx",
        default="docs/tez_ana_taslak_tr_frontmatter_ozetli.docx",
        help="Main thesis DOCX file to inspect.",
    )
    parser.add_argument(
        "--output-md",
        default="docs/tez_taslak_kalite_kontrol_raporu.md",
        help="Output Markdown quality report.",
    )
    parser.add_argument(
        "--output-json",
        default="docs/tez_taslak_kalite_kontrol_raporu.json",
        help="Output JSON quality report.",
    )
    args = parser.parse_args()

    docx_path = Path(args.docx)
    output_md = Path(args.output_md)
    output_json = Path(args.output_json)

    if not docx_path.exists():
        raise SystemExit(f"[ERROR] DOCX not found: {docx_path}")

    doc = Document(docx_path)

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    heading_positions = find_heading_positions(paragraphs, EXPECTED_HEADINGS)

    heading_checks = {
        h: {
            "found": heading_positions[h] is not None,
            "position": heading_positions[h],
        }
        for h in EXPECTED_HEADINGS
    }

    expected_file_checks = {
        path: {
            "exists": Path(path).exists(),
            "size": file_size(Path(path)),
        }
        for path in EXPECTED_FILES
    }

    artifact_checks = {
        path: {
            "exists": Path(path).exists(),
            "size": file_size(Path(path)),
        }
        for path in EXPECTED_ARTIFACTS
    }

    missing_headings = [h for h, v in heading_checks.items() if not v["found"]]
    missing_files = [p for p, v in expected_file_checks.items() if not v["exists"]]
    missing_artifacts = [p for p, v in artifact_checks.items() if not v["exists"]]

    # Front matter headings may contain chapter titles as manual TOC placeholders.
    # Therefore, chapter order should be checked by locating the last occurrence
    # of each real chapter heading, not the first occurrence found in the TOC.
    chapter_positions = {}
    for h in CHAPTER_HEADINGS:
        # Use exact normalized equality for real chapter headings to avoid matching TOC-like sentences.
        target = normalize_heading_text(h)
        positions = [
            i for i, text in enumerate(paragraphs)
            if normalize_heading_text(text) == target
        ]
        chapter_positions[h] = positions[-1] if positions else None

    frontmatter_positions = {}
    for h in FRONTMATTER_HEADINGS:
        positions = [
            i for i, text in enumerate(paragraphs)
            if heading_matches(text, h)
        ]
        frontmatter_positions[h] = positions[0] if positions else None

    # SAÜ FBE frontmatter order differs from the earlier simple draft order.
    # The generated frontmatter follows:
    # İçindekiler -> Tablo Listesi -> Şekil Listesi -> Özet -> Abstract -> numbered chapters.
    sau_fbe_order = [
        "İçindekiler",
        "Tablo Listesi",
        "Şekil Listesi",
        "Özet",
        "Abstract",
    ] + CHAPTER_HEADINGS

    heading_order_ok = True
    last_pos = -1

    for h in sau_fbe_order:
        if h in CHAPTER_HEADINGS:
            pos = chapter_positions.get(h)
        else:
            pos = frontmatter_positions.get(h)

        if pos is None or pos < last_pos:
            heading_order_ok = False
            break

        last_pos = pos

    summary = {
        "generated_at_utc": datetime.utcnow().isoformat(),
        "docx": str(docx_path),
        "docx_exists": docx_path.exists(),
        "docx_size": file_size(docx_path),
        "paragraph_count": len(doc.paragraphs),
        "non_empty_paragraph_count": len(paragraphs),
        "table_count": len(doc.tables),
        "inline_shape_count": len(doc.inline_shapes),
        "expected_table_count": 7,
        "expected_inline_shape_count": 4,
        "table_count_ok": len(doc.tables) >= 7,
        "inline_shape_count_ok": len(doc.inline_shapes) >= 4,
        "heading_order_ok": heading_order_ok,
        "missing_headings": missing_headings,
        "missing_files": missing_files,
        "missing_artifacts": missing_artifacts,
        "heading_checks": heading_checks,
        "frontmatter_positions": frontmatter_positions,
        "chapter_positions": chapter_positions,
        "expected_file_checks": expected_file_checks,
        "artifact_checks": artifact_checks,
    }

    # Recompute expected heading status with Turkish/case-insensitive matching
    output_json.parent.mkdir(parents=True, exist_ok=True)
    output_json.write_text(json.dumps(summary, indent=2, ensure_ascii=False), encoding="utf-8")

    lines = []
    lines.append("# Tez Taslak Kalite Kontrol Raporu")
    lines.append("")
    lines.append(f"- Generated at UTC: `{summary['generated_at_utc']}`")
    lines.append(f"- DOCX: `{summary['docx']}`")
    lines.append(f"- DOCX size: `{summary['docx_size']}`")
    lines.append("")
    lines.append("## 1. Genel DOCX Kontrolü")
    lines.append("")
    lines.append("| Kontrol | Değer | Durum |")
    lines.append("|---|---:|---|")
    lines.append(f"| Paragraph count | {summary['paragraph_count']} | ok |")
    lines.append(f"| Non-empty paragraph count | {summary['non_empty_paragraph_count']} | ok |")
    lines.append(f"| Table count | {summary['table_count']} | {'ok' if summary['table_count_ok'] else 'check'} |")
    lines.append(f"| Inline shape count | {summary['inline_shape_count']} | {'ok' if summary['inline_shape_count_ok'] else 'check'} |")
    lines.append(f"| Heading order | {summary['heading_order_ok']} | {'ok' if summary['heading_order_ok'] else 'check'} |")
    lines.append("")
    lines.append("## 2. Beklenen Başlık Kontrolü")
    lines.append("")
    lines.append("| Başlık | Durum | Pozisyon |")
    lines.append("|---|---|---:|")
    for h, v in heading_checks.items():
        lines.append(f"| {h} | {'found' if v['found'] else 'missing'} | {v['position'] if v['position'] is not None else '-'} |")
    lines.append("")
    lines.append("## 3. Beklenen Tez Dosyaları")
    lines.append("")
    lines.append("| Dosya | Durum | Boyut |")
    lines.append("|---|---|---:|")
    for path, v in expected_file_checks.items():
        lines.append(f"| `{path}` | {'exists' if v['exists'] else 'missing'} | {v['size']} |")
    lines.append("")
    lines.append("## 4. Tez Artifact Dosyaları")
    lines.append("")
    lines.append("| Artifact | Durum | Boyut |")
    lines.append("|---|---|---:|")
    for path, v in artifact_checks.items():
        lines.append(f"| `{path}` | {'exists' if v['exists'] else 'missing'} | {v['size']} |")
    lines.append("")
    lines.append("## 5. Özet Değerlendirme")
    lines.append("")
    if not missing_headings and not missing_files and not missing_artifacts and summary["table_count_ok"] and summary["inline_shape_count_ok"]:
        lines.append("Kalite kontrol sonucunda ana tez taslağında beklenen bölüm başlıkları, temel dosyalar, tablo artifact’leri ve şekil artifact’leri mevcut görünmektedir.")
    else:
        lines.append("Kalite kontrol sonucunda bazı eksik veya kontrol edilmesi gereken başlık/dosya/artifact bulunmuştur. Aşağıdaki listeler incelenmelidir.")
    lines.append("")
    lines.append(f"- Missing headings: `{missing_headings}`")
    lines.append(f"- Missing files: `{missing_files}`")
    lines.append(f"- Missing artifacts: `{missing_artifacts}`")
    lines.append("")

    try:
        all_texts
    except NameError:
        all_texts = [para.text.strip() for para in doc.paragraphs if para.text.strip()]

    append_sau_fbe_precheck_section(lines, doc, all_texts)

    output_md.parent.mkdir(parents=True, exist_ok=True)
    output_md.write_text("\n".join(lines), encoding="utf-8")

    print("[INFO] Quality report MD:", output_md)
    print("[INFO] Quality report JSON:", output_json)
    print("[INFO] Missing headings:", len(missing_headings))
    print("[INFO] Missing files:", len(missing_files))
    print("[INFO] Missing artifacts:", len(missing_artifacts))
    print("[INFO] Tables:", len(doc.tables))
    print("[INFO] Inline shapes:", len(doc.inline_shapes))


if __name__ == "__main__":
    main()
