#!/usr/bin/env python3
from __future__ import annotations

import argparse
from pathlib import Path
from docx import Document
import pandas as pd


def cell_text(cell):
    return " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip()).strip()


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--docx", default="docs/literature_review/source_files/Manuscript.docx")
    parser.add_argument("--output-dir", default="docs/literature_review/manuscript_alignment")
    args = parser.parse_args()

    docx_path = Path(args.docx)
    out_dir = Path(args.output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    if not docx_path.exists():
        raise SystemExit(f"[ERROR] Missing DOCX: {docx_path}")

    doc = Document(str(docx_path))

    inventory = []

    for ti, table in enumerate(doc.tables, start=1):
        rows = []
        max_cols = 0

        for row in table.rows:
            vals = [cell_text(c) for c in row.cells]
            max_cols = max(max_cols, len(vals))
            rows.append(vals)

        norm_rows = [r + [""] * (max_cols - len(r)) for r in rows]
        df = pd.DataFrame(norm_rows)
        csv_path = out_dir / f"manuscript_table_{ti:02d}.csv"
        df.to_csv(csv_path, index=False, header=False)

        preview = " | ".join(norm_rows[0][:min(max_cols, 6)]) if norm_rows else ""

        inventory.append({
            "table_no": ti,
            "rows": len(norm_rows),
            "cols": max_cols,
            "csv_path": str(csv_path),
            "first_row_preview": preview[:300],
        })

    inv = pd.DataFrame(inventory)
    inv_csv = out_dir / "manuscript_table_inventory.csv"
    inv_md = out_dir / "manuscript_table_inventory.md"

    inv.to_csv(inv_csv, index=False)

    md = []
    md.append("# Manuscript DOCX Table Inventory")
    md.append("")
    md.append(f"- DOCX: `{docx_path}`")
    md.append(f"- Table count: `{len(inventory)}`")
    md.append("")
    md.append("| Table | Rows | Cols | CSV | First row preview |")
    md.append("|---:|---:|---:|---|---|")
    for r in inventory:
        md.append(
            f"| {r['table_no']} | {r['rows']} | {r['cols']} | "
            f"`{r['csv_path']}` | {str(r['first_row_preview']).replace('|', '/')} |"
        )

    inv_md.write_text("\n".join(md), encoding="utf-8")

    print("[INFO] DOCX:", docx_path)
    print("[INFO] Tables:", len(inventory))
    print("[INFO] Inventory CSV:", inv_csv)
    print("[INFO] Inventory MD:", inv_md)


if __name__ == "__main__":
    main()
