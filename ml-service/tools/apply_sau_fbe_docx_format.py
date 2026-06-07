#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt


def set_run_font(run, font_name: str, font_size_pt: float) -> None:
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)


def apply_paragraph_font(paragraph, font_name: str, font_size_pt: float) -> None:
    for run in paragraph.runs:
        set_run_font(run, font_name, font_size_pt)


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Apply SAÜ FBE thesis template page and basic style settings to a DOCX file."
    )
    parser.add_argument("--input", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument(
        "--rules-json",
        default="docs/sau_fbe_format_rules.json",
    )
    parser.add_argument(
        "--body-font",
        default="Times New Roman",
    )
    parser.add_argument(
        "--body-size",
        type=float,
        default=12.0,
    )
    args = parser.parse_args()

    input_path = Path(args.input)
    output_path = Path(args.output)
    rules_path = Path(args.rules_json)

    if not input_path.exists():
        raise SystemExit(f"[ERROR] Input DOCX not found: {input_path}")
    if not rules_path.exists():
        raise SystemExit(f"[ERROR] Rules JSON not found: {rules_path}")

    rules = json.loads(rules_path.read_text(encoding="utf-8"))

    doc = Document(input_path)

    page = rules["page"]
    margins = rules["margins"]

    # Page setup
    for section in doc.sections:
        section.page_width = Cm(float(page["width_cm"]))
        section.page_height = Cm(float(page["height_cm"]))
        section.top_margin = Cm(float(margins["top_cm"]))
        section.bottom_margin = Cm(float(margins["bottom_cm"]))
        section.left_margin = Cm(float(margins["left_cm"]))
        section.right_margin = Cm(float(margins["right_cm"]))
        section.header_distance = Cm(float(margins["header_distance_cm"]))
        section.footer_distance = Cm(float(margins["footer_distance_cm"]))

    # Normal style
    normal = doc.styles["Normal"]
    normal.font.name = args.body_font
    normal.font.size = Pt(args.body_size)

    # Common heading styles. We keep headings bold but align font family.
    for style_name, size, bold in [
        ("Heading 1", 14, True),
        ("Heading 2", 13, True),
        ("Heading 3", 12, True),
        ("Heading 4", 12, True),
        ("Başlık 1", 14, True),
        ("Başlık 2", 13, True),
        ("Başlık 3", 12, True),
    ]:
        try:
            style = doc.styles[style_name]
            style.font.name = args.body_font
            style.font.size = Pt(size)
            style.font.bold = bold
        except Exception:
            pass

    # Apply body font to existing runs without destroying bold/italic.
    for paragraph in doc.paragraphs:
        apply_paragraph_font(paragraph, args.body_font, args.body_size)

    # Apply font to table cell paragraphs.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    apply_paragraph_font(paragraph, args.body_font, args.body_size)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    print("[INFO] Written:", output_path)


if __name__ == "__main__":
    main()
