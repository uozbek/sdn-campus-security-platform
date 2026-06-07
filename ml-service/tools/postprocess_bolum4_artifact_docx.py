#!/usr/bin/env python3
from argparse import ArgumentParser
from pathlib import Path
import re
from docx import Document


CAPTION_RE = re.compile(r"^(Tablo|Şekil)\s+(\d+(?:\.\d+)+)\.\s+(.+?)\s*$")


REPLACEMENTS = {
    "Bu aksiyon uzayı, binary model çıktısından daha zengindir. Bu nedenle deneysel değerlendirmede yalnızca exact action matching değil, security-compatible matching yaklaşımı da kullanılmıştır.":
    "Bu aksiyon uzayı, ikili model çıktısından daha zengin bir karar alanı sunmaktadır. Bu nedenle deneysel değerlendirmede yalnızca exact action matching değil, security-compatible matching yaklaşımı da kullanılmıştır.",

    "Runtime model çıktısı temel olarak binary bir sınıflandırma sunmaktadır: benign veya attack. Ancak SDN denetleyicisi tarafında bu çıktıların doğrudan kullanılması her zaman yeterli değildir. Özellikle TCP kontrol benzeri akışların, trafik bağlamı nedeniyle yüksek saldırı olasılığı alması mümkündür. Buna rağmen bu akışların yüksek hacimli UDP flood akışlarıyla aynı şekilde ele alınması doğru olmayacaktır.":
    "Runtime model çıktısı temel olarak ikili bir sınıflandırma sunmaktadır: normal ya da saldırı. Ancak SDN denetleyicisi tarafında bu çıktıların doğrudan kullanılması her zaman yeterli değildir. Özellikle TCP kontrol benzeri akışların, trafik bağlamı nedeniyle yüksek saldırı olasılığı alması mümkündür. Buna rağmen bu akışların yüksek hacimli UDP flood akışlarıyla aynı şekilde ele alınması doğru olmayacaktır.",

    "Bu durum exact action matching yerine security-compatible matching yaklaşımının neden daha anlamlı olduğunu göstermektedir. Çünkü binary model çıktısı DROP önerirken, denetleyici tekrarlanan yüksek güvenli saldırı davranışı nedeniyle aksiyonu quarantine seviyesine yükseltebilir. Dolayısıyla controller aksiyonunun model aksiyonu ile birebir aynı olması beklenmemelidir; önemli olan güvenlik açısından uyumlu bir önleme davranışının ortaya çıkmasıdır.":
    "Bu durum exact action matching yerine security-compatible matching yaklaşımının neden daha anlamlı olduğunu göstermektedir. Çünkü ikili model çıktısı saldırı yönünde güçlü bir sinyal üretirken, denetleyici tekrarlanan yüksek güvenli saldırı davranışı nedeniyle aksiyonu quarantine seviyesine yükseltebilir. Dolayısıyla controller aksiyonunun model aksiyonu ile birebir aynı olması beklenmemelidir; önemli olan güvenlik açısından uyumlu bir önleme davranışının ortaya çıkmasıdır.",

    "İkinci olarak, binary model çıktılarının SDN denetleyicisi tarafında daha zengin ve ağ bağlamına duyarlı bir politika yorumuna dönüştürülebileceği gösterilmiştir. Protocol-aware final policy katmanı TCP kontrol akışlarını, zararsız UDP trafiğini, zararlı UDP trafiğini ve quarantine sonrası gözlenen trafiği ayrı ayrı yorumlamaktadır.":
    "İkinci olarak, ikili model çıktılarının SDN denetleyicisi tarafında daha zengin ve ağ bağlamına duyarlı bir politika yorumuna dönüştürülebileceği gösterilmiştir. Protocol-aware final policy katmanı TCP kontrol akışlarını, zararsız UDP trafiğini, zararlı UDP trafiğini ve quarantine sonrası gözlenen trafiği ayrı ayrı yorumlamaktadır.",

    "Üçüncü olarak, exact action matching tek başına yeterli bir ölçüt değildir. Çünkü makine öğrenmesi modeli binary bir çıktı üretirken, SDN denetleyicisi rate-limit, drop ve quarantine gibi daha zengin bir aksiyon uzayına sahiptir. Bu nedenle security-compatible matching daha anlamlı bir değerlendirme yaklaşımıdır.":
    "Üçüncü olarak, exact action matching tek başına yeterli bir ölçüt değildir. Çünkü makine öğrenmesi modeli ikili bir sınıflandırma çıktısı üretirken, SDN denetleyicisi rate-limit, drop ve quarantine gibi daha zengin bir aksiyon uzayına sahiptir. Bu nedenle security-compatible matching daha anlamlı bir değerlendirme yaklaşımıdır.",
}


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def replace_paragraph_text(paragraph, new_text):
    # Basit ama güvenli yöntem: paragraf run'larını temizle, ilk run'a yeni metni yaz.
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)


def main():
    ap = ArgumentParser()
    ap.add_argument("--docx", required=True)
    ap.add_argument("--out", required=True)
    args = ap.parse_args()

    in_path = Path(args.docx)
    out_path = Path(args.out)

    doc = Document(in_path)

    removed_source_paths = []
    removed_duplicate_captions = []
    replaced_texts = []

    seen_captions = set()

    for p in list(doc.paragraphs):
        text = p.text.strip()

        # Kaynak path ve insertion marker temizliği
        if (
            text.startswith("Kaynak: tables/")
            or text.startswith("Kaynak: figures/")
            or "TABLO EKLEME" in text
            or "ŞEKİL EKLEME" in text
        ):
            removed_source_paths.append(text)
            remove_paragraph(p)
            continue

        # Aynı caption ikinci kez gelirse sil.
        m = CAPTION_RE.match(text)
        if m:
            key = (m.group(1), m.group(2), m.group(3).strip())
            if key in seen_captions:
                removed_duplicate_captions.append(text)
                remove_paragraph(p)
                continue
            seen_captions.add(key)

        # Eski akademik olmayan ifadeleri değiştir.
        if text in REPLACEMENTS:
            replace_paragraph_text(p, REPLACEMENTS[text])
            replaced_texts.append(text)

    doc.save(out_path)

    print(f"[INFO] Input : {in_path}")
    print(f"[INFO] Output: {out_path}")
    print(f"[INFO] Removed source/marker paragraphs: {len(removed_source_paths)}")
    print(f"[INFO] Removed duplicate captions: {len(removed_duplicate_captions)}")
    print(f"[INFO] Replaced text paragraphs: {len(replaced_texts)}")

    for x in removed_duplicate_captions[:30]:
        print(f"[REMOVED_DUP_CAPTION] {x}")

    for x in replaced_texts:
        print(f"[REPLACED] {x[:120]}")


if __name__ == "__main__":
    main()
