#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
from datetime import datetime
from pathlib import Path

from docx import Document


KEYWORDS = [
    "SAKARYA ÜNİVERSİTESİ",
    "FEN BİLİMLERİ ENSTİTÜSÜ",
    "DOKTORA TEZİ",
    "YÜKSEK LİSANS TEZİ",
    "TEZ ONAY",
    "ONAY",
    "BEYAN",
    "TEŞEKKÜR",
    "İÇİNDEKİLER",
    "KISALTMALAR",
    "SİMGELER",
    "TABLO LİSTESİ",
    "ŞEKİL LİSTESİ",
    "ÖZET",
    "SUMMARY",
    "ABSTRACT",
]


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--template-docx",
        default="docs/thesis_template_sources/tez_hazirlama_sablonu.docx",
    )
    parser.add_argument(
        "--output-md",
        default="docs/template_frontmatter_structure_report.md",
    )
    parser.add_argument(
        "--output-json",
        default="docs/template_frontmatter_structure_report.json",
    )
    args = parser.parse_args()

    p = Path(args.template_docx)
    if not p.exists():
        raise SystemExit(f"[ERROR] Template not found: {p}")

    doc = Document(p)

    non_empty = []
    keyword_hits = []

    for i, para in enumerate(doc.paragraphs, start=1):
        txt = para.text.strip()
        if not txt:
            continue
        rec = {
            "paragraph_index": i,
            "style": para.style.name,
            "text": txt,
        }
        non_empty.append(rec)

        upper = txt.upper()
        for kw in KEYWORDS:
            if kw in upper:
                keyword_hits.append({
                    "paragraph_index": i,
                    "style": para.style.name,
                    "keyword": kw,
                    "text": txt,
                })

    result = {
        "generated_at_utc": datetime.utcnow().isoformat(),
        "template_docx": str(p),
        "paragraph_count": len(doc.paragraphs),
        "non_empty_count": len(non_empty),
        "table_count": len(doc.tables),
        "inline_shape_count": len(doc.inline_shapes),
        "keyword_hits": keyword_hits,
        "first_220_non_empty": non_empty[:220],
    }

    Path(args.output_json).write_text(json.dumps(result, indent=2, ensure_ascii=False), encoding="utf-8")

    md = []
    md.append("# Tez Şablonu Ön Sayfa Yapısı Raporu")
    md.append("")
    md.append(f"- Generated at UTC: `{result['generated_at_utc']}`")
    md.append(f"- Template DOCX: `{p}`")
    md.append(f"- Paragraphs: `{result['paragraph_count']}`")
    md.append(f"- Non-empty paragraphs: `{result['non_empty_count']}`")
    md.append(f"- Tables: `{result['table_count']}`")
    md.append(f"- Inline shapes: `{result['inline_shape_count']}`")
    md.append("")
    md.append("## 1. Anahtar Başlık Eşleşmeleri")
    md.append("")
    if keyword_hits:
        md.append("| P | Style | Keyword | Text |")
        md.append("|---:|---|---|---|")
        for h in keyword_hits:
            txt = h["text"].replace("|", "\\|")
            md.append(f"| {h['paragraph_index']} | {h['style']} | {h['keyword']} | {txt[:180]} |")
    else:
        md.append("_Anahtar başlık bulunamadı._")

    md.append("")
    md.append("## 2. İlk 220 Dolu Paragraf")
    md.append("")
    md.append("| P | Style | Text |")
    md.append("|---:|---|---|")
    for item in non_empty[:220]:
        txt = item["text"].replace("|", "\\|")
        md.append(f"| {item['paragraph_index']} | {item['style']} | {txt[:220]} |")

    Path(args.output_md).write_text("\n".join(md), encoding="utf-8")

    print("[INFO] MD:", args.output_md)
    print("[INFO] JSON:", args.output_json)


if __name__ == "__main__":
    main()
