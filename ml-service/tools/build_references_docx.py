#!/usr/bin/env python3
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--references-md", default="docs/references_apa_like.md")
    parser.add_argument("--output", default="docs/bolum_kaynakca_tr.docx")
    args = parser.parse_args()

    md = Path(args.references_md)
    out = Path(args.output)

    refs = []
    for line in md.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if line.startswith("- "):
            refs.append(line[2:].strip())

    if not refs:
        raise SystemExit(f"[ERROR] No references found in {md}")

    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(4.0)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    h = doc.add_paragraph("Kaynakça")
    h.style = styles["Heading 1"] if "Heading 1" in styles else normal
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.bold = True

    for ref in refs:
        p = doc.add_paragraph(ref)
        p.style = normal
        p.paragraph_format.first_line_indent = Cm(-0.75)
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.0

        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)

    print("[INFO] Written:", out)
    print("[INFO] Reference count:", len(refs))


if __name__ == "__main__":
    main()
