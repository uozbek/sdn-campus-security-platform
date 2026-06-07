#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from datetime import datetime

from docx import Document


CAPTION_RE = re.compile(
    r"^(Tablo|Şekil)\s+([0-9]+(?:\.[0-9]+|\.x|\.X)*)\s*[\.\-:]?\s+(.+)$"
)

TECH_REF_RE = re.compile(r"\[(?:BIB|LR|MAN)\d+\]")
APA_INLINE_RE = re.compile(
    r"\(([A-ZÇĞİÖŞÜ][A-Za-zÇĞİÖŞÜçğıöşü\- ]+(?:\s+vd\.)?|[A-ZÇĞİÖŞÜ][A-Za-zÇĞİÖŞÜçğıöşü\- ]+\s+ve\s+[A-ZÇĞİÖŞÜ][A-Za-zÇĞİÖŞÜçğıöşü\- ]+),\s*(19|20)\d{2}\)"
)
APA_REF_RE = re.compile(
    r"^[A-ZÇĞİÖŞÜ][^\n]{2,240}\((?:19|20)\d{2}\)\."
)


def iter_paragraph_texts(doc: Document):
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            yield t

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = p.text.strip()
                    if t:
                        yield t


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--docx", required=True)
    parser.add_argument("--output-md", required=True)
    parser.add_argument("--output-json", required=True)
    args = parser.parse_args()

    docx_path = Path(args.docx)
    if not docx_path.exists():
        raise SystemExit(f"[ERROR] DOCX not found: {docx_path}")

    doc = Document(docx_path)

    paragraph_texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    all_texts = list(iter_paragraph_texts(doc))

    table_captions = []
    figure_captions = []

    for idx, text in enumerate(paragraph_texts, start=1):
        m = CAPTION_RE.match(text)
        if not m:
            continue
        kind, number, title = m.groups()
        item = {
            "paragraph": idx,
            "kind": kind,
            "number": number,
            "title": title,
            "text": text,
        }
        if kind == "Tablo":
            table_captions.append(item)
        elif kind == "Şekil":
            figure_captions.append(item)

    joined = "\n".join(all_texts)

    technical_reference_marker_count = len(TECH_REF_RE.findall(joined))
    apa_inline_citation_count = len(APA_INLINE_RE.findall(joined))

    apa_reference_like = []
    kaynakca_seen = False
    for idx, text in enumerate(paragraph_texts, start=1):
        if text.lower() in {"kaynakça", "kaynaklar", "references", "bibliography"}:
            kaynakca_seen = True
            continue
        if kaynakca_seen and APA_REF_RE.match(text):
            apa_reference_like.append({"paragraph": idx, "text": text})

    table_count = len(doc.tables)
    shape_count = len(doc.inline_shapes)

    checks = []

    checks.append((
        "Word table object count",
        ">= 1",
        table_count,
        "ok" if table_count >= 1 else "check",
    ))

    checks.append((
        "Word inline shape count",
        ">= 1",
        shape_count,
        "ok" if shape_count >= 1 else "check",
    ))

    # Her Word tablosunun caption olması şart değil; bazı tablolar layout/artefact olabilir.
    checks.append((
        "Detected table caption count",
        ">= 1 and manually consistent with table_count",
        len(table_captions),
        "ok" if len(table_captions) >= 1 else "check",
    ))

    checks.append((
        "Detected figure caption count",
        ">= inline shape count or manually justified",
        len(figure_captions),
        "ok" if len(figure_captions) >= shape_count else "check",
    ))

    reference_signal_count = (
        technical_reference_marker_count
        + apa_inline_citation_count
        + len(apa_reference_like)
    )

    checks.append((
        "Reference signal count",
        ">= 1 technical marker, APA inline citation, or APA-like bibliography entry",
        f"technical_markers={technical_reference_marker_count}; apa_inline={apa_inline_citation_count}; apa_reference_like={len(apa_reference_like)}",
        "ok" if reference_signal_count >= 1 else "check",
    ))

    kaynakca_present = any(t.lower() == "kaynakça" for t in paragraph_texts)
    checks.append((
        "Kaynakça heading",
        "present",
        "present" if kaynakca_present else "missing",
        "ok" if kaynakca_present else "check",
    ))

    ok_count = sum(1 for *_, status in checks if status == "ok")
    check_count = sum(1 for *_, status in checks if status != "ok")

    summary = {
        "generated_at_utc": datetime.utcnow().isoformat(),
        "docx": str(docx_path),
        "ok": ok_count,
        "check": check_count,
        "total": len(checks),
        "word_table_count": table_count,
        "inline_shape_count": shape_count,
        "detected_table_caption_count": len(table_captions),
        "detected_figure_caption_count": len(figure_captions),
        "technical_reference_marker_count": technical_reference_marker_count,
        "apa_inline_citation_count": apa_inline_citation_count,
        "apa_reference_like_count": len(apa_reference_like),
        "checks": [
            {
                "name": c[0],
                "expected": c[1],
                "current": c[2],
                "status": c[3],
            }
            for c in checks
        ],
        "table_captions": table_captions,
        "figure_captions": figure_captions,
        "apa_reference_like_examples": apa_reference_like[:20],
    }

    md = []
    md.append("# Tez Tablo/Şekil Başlığı ve Kaynakça Ön Kontrol Raporu")
    md.append("")
    md.append(f"- Generated at UTC: `{summary['generated_at_utc']}`")
    md.append(f"- DOCX: `{docx_path}`")
    md.append("")
    md.append("## 1. Özet")
    md.append("")
    md.append(f"- OK: `{ok_count}`")
    md.append(f"- Check: `{check_count}`")
    md.append(f"- Total: `{len(checks)}`")
    md.append(f"- Word table count: `{table_count}`")
    md.append(f"- Inline shape count: `{shape_count}`")
    md.append(f"- Detected table captions: `{len(table_captions)}`")
    md.append(f"- Detected figure captions: `{len(figure_captions)}`")
    md.append(f"- Technical reference marker count: `{technical_reference_marker_count}`")
    md.append(f"- APA inline citation count: `{apa_inline_citation_count}`")
    md.append(f"- APA reference-like count: `{len(apa_reference_like)}`")
    md.append("")
    md.append("## 2. Kontroller")
    md.append("")
    md.append("| Kontrol | Beklenen | Mevcut | Durum |")
    md.append("|---|---|---:|---|")
    for name, expected, current, status in checks:
        md.append(f"| {name} | {expected} | {current} | {status} |")

    md.append("")
    md.append("## 3. Bulunan Tablo Başlığı Örnekleri")
    md.append("")
    if table_captions:
        for c in table_captions[:30]:
            md.append(f"- P{c['paragraph']}: {c['text']}")
    else:
        md.append("Tablo başlığı bulunamadı.")

    md.append("")
    md.append("## 4. Bulunan Şekil Başlığı Örnekleri")
    md.append("")
    if figure_captions:
        for c in figure_captions[:30]:
            md.append(f"- P{c['paragraph']}: {c['text']}")
    else:
        md.append("Şekil başlığı bulunamadı.")

    md.append("")
    md.append("## 5. Kaynakça / Referans Sinyalleri")
    md.append("")
    md.append(f"- Teknik marker sayısı: `{technical_reference_marker_count}`")
    md.append(f"- APA metin içi atıf sinyali sayısı: `{apa_inline_citation_count}`")
    md.append(f"- Kaynakça bölümündeki APA-benzeri kaynak satırı sayısı: `{len(apa_reference_like)}`")
    md.append("")
    md.append("## 6. Not")
    md.append("")
    md.append(
        "Bu rapor otomatik ön kontroldür. Tablo/şekil başlıklarının gerçek Word caption alanı olup olmadığı "
        "ve Word Tablo/Şekil Listesi ile uyumu nihai aşamada Microsoft Word üzerinde ayrıca kontrol edilmelidir."
    )

    Path(args.output_md).write_text("\n".join(md), encoding="utf-8")
    Path(args.output_json).write_text(
        json.dumps(summary, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    print("[INFO] MD:", args.output_md)
    print("[INFO] JSON:", args.output_json)
    print("[INFO] Summary:", {"ok": ok_count, "check": check_count, "total": len(checks)})
    print("[INFO] Word tables:", table_count)
    print("[INFO] Inline shapes:", shape_count)
    print("[INFO] Table captions:", len(table_captions))
    print("[INFO] Figure captions:", len(figure_captions))


if __name__ == "__main__":
    main()
