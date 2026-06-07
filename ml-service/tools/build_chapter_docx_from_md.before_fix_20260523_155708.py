#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def clean_line(s: str) -> str:
    return s.rstrip()


def add_formatted_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()

    # Markdown biçim işaretlerini düz metne çevir.
    text = text.replace("**", "").replace("__", "")
    text = re.sub(r"`([^`]+)`", r"\1", text)

    # Markdown linklerini "metin (url)" biçimine çevir.
    # Örnek: [OpenFlow](https://example.com) -> OpenFlow (https://example.com)
    text = re.sub(r"\$begin:math:display$\(\[\^\$end:math:display$]+)\]\$begin:math:text$\(\[\^\)\]\+\)\$end:math:text$", r"\1 (\2)", text)

    # Fazla boşlukları sadeleştir.
    text = re.sub(r"\s+", " ", text).strip()

    p.add_run(text)
    return p


def apply_base_formatting(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(4.0)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        if style_name in styles:
            st = styles[style_name]
            st.font.name = "Times New Roman"
            st.font.bold = True
            if style_name == "Heading 1":
                st.font.size = Pt(14)
            else:
                st.font.size = Pt(12)


def md_to_docx(md_path: Path, out_path: Path):
    doc = Document()
    apply_base_formatting(doc)

    lines = md_path.read_text(encoding="utf-8").splitlines()

    in_table = False
    table_buffer = []

    def flush_table():
        nonlocal table_buffer
        if not table_buffer:
            return

        rows = []
        for line in table_buffer:
            if re.match(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", line):
                continue
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            if cells:
                rows.append(cells)

        if rows:
            max_cols = max(len(r) for r in rows)
            table = doc.add_table(rows=0, cols=max_cols)
            table.style = "Table Grid"

            for r_idx, row in enumerate(rows):
                cells = table.add_row().cells
                for i in range(max_cols):
                    cells[i].text = row[i] if i < len(row) else ""
                    for para in cells[i].paragraphs:
                        for run in para.runs:
                            run.font.name = "Times New Roman"
                            run.font.size = Pt(10)
                        if r_idx == 0:
                            for run in para.runs:
                                run.bold = True

        table_buffer = []

    for raw in lines:
        line = clean_line(raw)

        if not line.strip():
            flush_table()
            continue

        if line.lstrip().startswith("<!--"):
            continue

        if line.strip().startswith("|") and "|" in line.strip()[1:]:
            table_buffer.append(line)
            continue
        else:
            flush_table()

        if line.startswith("# "):
            p = doc.add_paragraph(line[2:].strip(), style="Heading 1")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith("## "):
            p = doc.add_paragraph(line[3:].strip(), style="Heading 2")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith("### "):
            p = doc.add_paragraph(line[4:].strip(), style="Heading 3")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif re.match(r"^\s*[-*]\s+", line):
            p = add_formatted_paragraph(doc, re.sub(r"^\s*[-*]\s+", "", line))
            p.style = doc.styles["Normal"]
        else:
            add_formatted_paragraph(doc, line)

    flush_table()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--md", required=True)
    parser.add_argument("--out", required=True)
    args = parser.parse_args()

    md_path = Path(args.md)
    out_path = Path(args.out)

    if not md_path.exists():
        raise SystemExit(f"[ERROR] Missing MD: {md_path}")

    md_to_docx(md_path, out_path)
    print("[INFO] DOCX:", out_path)


if __name__ == "__main__":
    main()
