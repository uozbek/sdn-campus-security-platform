#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def clean_md(text: str) -> str:
    return (
        text.replace("`", "")
        .replace("**", "")
        .replace("__", "")
        .strip()
    )


def configure_doc(doc: Document) -> None:
    styles = doc.styles

    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Times New Roman"
        styles[style_name].font.bold = True

    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 3"].font.size = Pt(12)

    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)


def add_line(doc: Document, line: str) -> None:
    stripped = line.strip()

    if not stripped:
        return

    if stripped.startswith("# "):
        doc.add_heading(clean_md(stripped[2:]), level=1)
        return

    if stripped.startswith("## "):
        doc.add_heading(clean_md(stripped[3:]), level=2)
        return

    if stripped.startswith("### "):
        doc.add_heading(clean_md(stripped[4:]), level=3)
        return

    if stripped.startswith("- "):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(clean_md(stripped[2:]))
        return

    if re.match(r"^\d+\.\s+", stripped):
        p = doc.add_paragraph(style="List Number")
        p.add_run(clean_md(re.sub(r"^\d+\.\s+", "", stripped)))
        return

    if stripped == "---":
        doc.add_page_break()
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(clean_md(stripped))


def build_docx(markdown_path: Path, output_path: Path) -> None:
    doc = Document()
    configure_doc(doc)

    for line in markdown_path.read_text(encoding="utf-8").splitlines():
        add_line(doc, line)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--markdown", required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()

    build_docx(Path(args.markdown), Path(args.output))
    print(f"[INFO] Written: {args.output}")


if __name__ == "__main__":
    main()
