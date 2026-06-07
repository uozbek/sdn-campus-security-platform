#!/usr/bin/env python3
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document


TR_ABSTRACT = (
    "Bu tez çalışmasında, kampüs ağı benzeri yazılım tanımlı ağ ortamlarında DDoS saldırılarının "
    "makine öğrenmesi destekli olarak tespit edilmesi ve SDN denetleyicisi üzerinden aktif önleme "
    "aksiyonlarına dönüştürülmesi amaçlanmıştır. Bu amaçla Mininet, Open vSwitch, Ryu denetleyicisi, "
    "FastAPI tabanlı makine öğrenmesi servisi ve PCAP tabanlı çalışma zamanı özellik çıkarımı bileşenlerinden "
    "oluşan uçtan uca bir IDS/IPS prototipi geliştirilmiştir. Çalışmada CIC-DDoS2019 veri kümesi ailesiyle "
    "uyumlu seçilmiş özellikler kullanılmış ve Final XGBoost Top-20 modeli çalışma zamanı tespit bileşeni "
    "olarak sisteme entegre edilmiştir. Deneysel doğrulama sürecinde benign TCP/UDP trafik ile yüksek hacimli "
    "zararlı UDP trafik aynı SDN ortamında üretilmiş; model tahminleri, Ryu denetleyici politika kararları "
    "ve OpenFlow tabanlı önleme kayıtları birlikte değerlendirilmiştir. Port-aware ve protocol-aware canonical "
    "runtime validation koşusu olan run_05_port_aware_repeat_validation sonucunda benign/control akışların "
    "korunduğu, malicious UDP akışların yüksek saldırı olasılığı ile tespit edildiği ve rate-limit, drop ve "
    "quarantine aksiyonlarının gözlemlendiği görülmüştür. Canonical aggregate sonuçları, run_03_aligned_clean "
    "ve run_05_port_aware_repeat_validation koşularında model-controller eşleşmesinin ve güvenlik uyumlu "
    "aksiyonların tekrarlı biçimde üretilebildiğini göstermektedir. Elde edilen bulgular, önerilen mimarinin "
    "yalnızca çevrimdışı bir sınıflandırma yaklaşımı değil, çalışma zamanı SDN tabanlı aktif IDS/IPS prototipi "
    "olarak da uygulanabilir olduğunu ortaya koymaktadır."
)

TR_KEYWORDS = (
    "Anahtar Kelimeler: Yazılım Tanımlı Ağlar, DDoS, Makine Öğrenmesi, XGBoost, IDS, IPS, Ryu, OpenFlow, "
    "Mininet, Çalışma Zamanı Doğrulama"
)

EN_ABSTRACT = (
    "This dissertation investigates machine-learning-assisted DDoS detection and active prevention in "
    "software-defined campus-like network environments. To this end, an end-to-end IDS/IPS prototype was "
    "developed by integrating Mininet, Open vSwitch, a Ryu-based SDN controller, a FastAPI-based machine "
    "learning inference service, and a PCAP-based runtime feature extraction pipeline. The proposed system "
    "uses selected features compatible with the CIC-DDoS2019 dataset family, and the Final XGBoost Top-20 "
    "model is integrated as the runtime detection component. During experimental validation, benign TCP/UDP "
    "traffic and high-volume malicious UDP traffic were generated in the same SDN environment. Runtime model "
    "predictions, Ryu controller policy decisions, and OpenFlow-based enforcement logs were jointly analyzed. "
    "In the port-aware and protocol-aware canonical runtime validation run, namely "
    "run_05_port_aware_repeat_validation, benign/control flows were preserved, malicious UDP flows were "
    "detected with high attack probability, and rate-limit, drop, and quarantine actions were observed. "
    "The canonical aggregate results from run_03_aligned_clean and run_05_port_aware_repeat_validation further "
    "show that model-controller alignment and security-compatible enforcement actions can be reproduced across "
    "validated runtime experiments. The findings indicate that the proposed architecture is not limited to "
    "offline classification, but can also operate as an active SDN-based runtime IDS/IPS prototype."
)

EN_KEYWORDS = (
    "Keywords: Software-Defined Networking, DDoS, Machine Learning, XGBoost, IDS, IPS, Ryu, OpenFlow, "
    "Mininet, Runtime Validation"
)


def replace_section_content(doc: Document, heading: str, new_paragraphs: list[str], stop_headings: set[str]) -> bool:
    paras = doc.paragraphs
    start_idx = None

    for i, p in enumerate(paras):
        if p.text.strip() == heading:
            start_idx = i
            break

    if start_idx is None:
        return False

    end_idx = len(paras)
    for j in range(start_idx + 1, len(paras)):
        txt = paras[j].text.strip()
        if txt in stop_headings:
            end_idx = j
            break

    # Remove existing paragraphs between heading and next stop heading.
    for k in range(end_idx - 1, start_idx, -1):
        elem = paras[k]._element
        elem.getparent().remove(elem)

    # Insert new paragraphs after heading, preserving order.
    anchor = paras[start_idx]._element
    parent = anchor.getparent()
    pos = parent.index(anchor) + 1

    for text in new_paragraphs:
        new_p = doc.add_paragraph(text)
        elem = new_p._element
        elem.getparent().remove(elem)
        parent.insert(pos, elem)
        pos += 1

    return True


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()

    input_path = Path(args.input)
    output_path = Path(args.output)

    if not input_path.exists():
        raise SystemExit(f"[ERROR] Input not found: {input_path}")

    doc = Document(input_path)

    tr_ok = replace_section_content(
        doc,
        "Özet",
        [TR_ABSTRACT, TR_KEYWORDS],
        {"Abstract", "İçindekiler"},
    )

    en_ok = replace_section_content(
        doc,
        "Abstract",
        [EN_ABSTRACT, EN_KEYWORDS],
        {"İçindekiler", "Tablo Listesi"},
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    print("[INFO] Written:", output_path)
    print("[INFO] Updated Özet:", tr_ok)
    print("[INFO] Updated Abstract:", en_ok)


if __name__ == "__main__":
    main()
