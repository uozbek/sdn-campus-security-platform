#!/usr/bin/env python3
from pathlib import Path
import argparse
import csv

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


TABLES = [
    ("Tablo 1. Canonical tekrarlı çalışma zamanı doğrulama sonuçları", "tables/table_canonical_runtime_validation_summary.csv"),
    ("Tablo 2. Denetleyici taraflı IDS/IPS politika aksiyon dağılımı", "tables/table_controller_action_distribution.csv"),
    ("Tablo 3. Final XGBoost Top-20 modelinin runtime tahmin dağılımı", "tables/table_final_top20_prediction_distribution.csv"),
    ("Tablo 4. Protocol-aware final policy dağılımı", "tables/table_protocol_aware_final_policy_distribution.csv"),
    ("Tablo 5. SDN denetleyicisi önleme aksiyon özeti", "tables/table_enforcement_action_summary.csv"),
    ("Tablo 6. Flow-level model-controller karşılaştırması", "tables/table_flow_level_model_controller_comparison.csv"),
    ("Tablo 7. Flow-level saldırı olasılığı ve protocol-aware final action özeti", "tables/table_protocol_aware_attack_probability_summary.csv"),
]

FIGURES = [
    ("Şekil 1. Denetleyici taraflı politika aksiyon dağılımı", "figures/fig_controller_action_distribution.png"),
    ("Şekil 2. Runtime BENIGN/ATTACK tahmin dağılımı", "figures/fig_final_top20_prediction_distribution.png"),
    ("Şekil 3. Protocol-aware final policy aksiyon dağılımı", "figures/fig_protocol_aware_final_policy_distribution.png"),
    ("Şekil 4. SDN denetleyicisi önleme aksiyon özeti", "figures/fig_enforcement_action_summary.png"),
]


def clean_md(line: str) -> str:
    return line.replace("`", "").replace("**", "").strip()


def add_markdown_text(doc: Document, md_path: Path) -> None:
    for raw in md_path.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line:
            continue

        # Marker/placeholders are not needed in this appendix-based version.
        if "EKLEME NOKTASI" in line or line.startswith("> Dosya:") or line.startswith("> Başlık:"):
            continue

        if line.startswith("# "):
            doc.add_heading(clean_md(line[2:]), level=1)
        elif line.startswith("## "):
            doc.add_heading(clean_md(line[3:]), level=2)
        elif line.startswith("### "):
            doc.add_heading(clean_md(line[4:]), level=3)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(clean_md(line[2:]))
        elif line == "---":
            doc.add_page_break()
        elif line.startswith("> "):
            p = doc.add_paragraph()
            run = p.add_run(clean_md(line[2:]))
            run.italic = True
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run(clean_md(line))


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10)


def add_csv_table(doc: Document, csv_path: Path, max_rows: int = 15) -> None:
    if not csv_path.exists():
        doc.add_paragraph(f"[UYARI: Tablo dosyası bulunamadı: {csv_path}]")
        return

    with csv_path.open("r", encoding="utf-8") as f:
        rows = list(csv.reader(f))

    if not rows:
        doc.add_paragraph(f"[UYARI: Tablo boş: {csv_path}]")
        return

    truncated = False
    if len(rows) > max_rows + 1:
        rows = rows[: max_rows + 1]
        truncated = True

    col_count = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row in enumerate(rows):
        for j in range(col_count):
            value = row[j] if j < len(row) else ""
            cell = table.cell(i, j)
            cell.text = str(value)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(7)
                    if i == 0:
                        run.bold = True

    if truncated:
        p = doc.add_paragraph("Not: Tablo Word düzeni için kısaltılmıştır; tam içerik CSV dosyasında yer almaktadır.")
        p.runs[0].italic = True


def add_figure(doc: Document, image_path: Path, caption: str) -> None:
    if not image_path.exists():
        doc.add_paragraph(f"[UYARI: Şekil dosyası bulunamadı: {image_path}]")
        return

    doc.add_picture(str(image_path), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, caption)


def configure(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)

    for s in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[s].font.name = "Times New Roman"
        styles[s].font.bold = True

    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--markdown", required=True)
    ap.add_argument("--artifact-dir", required=True)
    ap.add_argument("--output", required=True)
    args = ap.parse_args()

    md_path = Path(args.markdown)
    artifact_dir = Path(args.artifact_dir)
    output = Path(args.output)

    doc = Document()
    configure(doc)

    add_markdown_text(doc, md_path)

    doc.add_page_break()
    doc.add_heading("Ek: Tez Tablo ve Şekilleri", level=1)

    doc.add_heading("Tablolar", level=2)
    for caption, rel in TABLES:
        add_caption(doc, caption)
        add_csv_table(doc, artifact_dir / rel)
        doc.add_paragraph()

    doc.add_page_break()
    doc.add_heading("Şekiller", level=2)
    for caption, rel in FIGURES:
        add_figure(doc, artifact_dir / rel, caption)
        doc.add_paragraph()

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(f"[INFO] Written: {output}")


if __name__ == "__main__":
    main()
