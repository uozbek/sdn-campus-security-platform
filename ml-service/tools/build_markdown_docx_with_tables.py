#!/usr/bin/env python3
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.shared import Pt


def is_table_separator(line: str) -> bool:
    line = line.strip()
    if not line.startswith("|") or not line.endswith("|"):
        return False
    cells = [c.strip() for c in line.strip("|").split("|")]
    return all(c and set(c) <= set("-:") for c in cells)


def is_table_line(line: str) -> bool:
    line = line.strip()
    return line.startswith("|") and line.endswith("|") and "|" in line[1:-1]


def parse_table_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def add_markdown_table(doc: Document, table_lines: list[str]) -> None:
    if len(table_lines) < 2:
        return

    header = parse_table_row(table_lines[0])
    body_lines = [ln for ln in table_lines[2:] if is_table_line(ln)]
    rows = [parse_table_row(ln) for ln in body_lines]

    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, val in enumerate(header):
        hdr_cells[i].text = val

    for row in rows:
        cells = table.add_row().cells
        for i in range(len(header)):
            cells[i].text = row[i] if i < len(row) else ""

    doc.add_paragraph("")


def add_paragraph_or_heading(doc: Document, line: str) -> None:
    stripped = line.strip()

    if not stripped:
        doc.add_paragraph("")
        return

    if stripped.startswith("# "):
        doc.add_heading(stripped[2:].strip(), level=1)
    elif stripped.startswith("## "):
        doc.add_heading(stripped[3:].strip(), level=2)
    elif stripped.startswith("### "):
        doc.add_heading(stripped[4:].strip(), level=3)
    elif stripped.startswith("#### "):
        doc.add_heading(stripped[5:].strip(), level=4)
    elif stripped.startswith("- "):
        doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
    elif stripped.startswith("* "):
        doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
    else:
        # Simple bold cleanup for captions like **Tablo ...**
        if stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
            p = doc.add_paragraph()
            r = p.add_run(stripped.strip("*"))
            r.bold = True
        else:
            doc.add_paragraph(stripped)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--markdown", required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()

    md_path = Path(args.markdown)
    out_path = Path(args.output)

    if not md_path.exists():
        raise SystemExit(f"[ERROR] Markdown file not found: {md_path}")

    lines = md_path.read_text(encoding="utf-8").splitlines()

    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i]

        if is_table_line(line):
            table_lines = []
            while i < len(lines) and is_table_line(lines[i]):
                table_lines.append(lines[i])
                i += 1

            if len(table_lines) >= 2 and is_table_separator(table_lines[1]):
                add_markdown_table(doc, table_lines)
            else:
                for ln in table_lines:
                    add_paragraph_or_heading(doc, ln)
            continue

        add_paragraph_or_heading(doc, line)
        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)

    print("[INFO] Written:", out_path)


if __name__ == "__main__":
    main()
