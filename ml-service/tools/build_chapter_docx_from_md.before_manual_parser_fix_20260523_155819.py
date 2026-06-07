#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def normalize_inline_markdown(text: str) -> str:
    """Convert simple Markdown inline syntax to plain text."""
    text = text.replace("**", "").replace("__", "")
    text = re.sub(r"`([^`]+)`", r"\1", text)

    # Markdown link: [label](url) -> label (url)
    text = re.sub(r"$begin:math:display$\(\[\^$end:math:display$]+)\]$begin:math:text$\(\[\^\)\]\+\)$end:math:text$", r"\1 (\2)", text)

    # Remove HTML comments
    text = re.sub(r"<!--.*?-->", "", text)

    text = re.sub(r"\s+", " ", text).strip()
    return text


def apply_base_formatting(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(4.0)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for name, size in [("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 12)]:
        if name in doc.styles:
            st = doc.styles[name]
            st.font.name = "Times New Roman"
            st.font.size = Pt(size)
            st.font.bold = True


def add_plain_paragraph(doc: Document, text: str):
    text = normalize_inline_markdown(text)
    if not text:
        return None
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return p


def add_markdown_table(doc: Document, table_lines: list[str]) -> None:
    rows = []
    for line in table_lines:
        s = line.strip()
        if re.match(r"^\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", s):
            continue
        cells = [normalize_inline_markdown(c.strip()) for c in s.strip("|").split("|")]
        if cells:
            rows.append(cells)

    if not rows:
        return

    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=max_cols)
    table.style = "Table Grid"

    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i in range(max_cols):
            cells[i].text = row[i] if i < len(row) else ""
            for para in cells[i].paragraphs:
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)
                    if r_idx == 0:
                        run.bold = True


def md_to_docx(md_path: Path, out_path: Path) -> None:
    doc = Document()
    apply_base_formatting(doc)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    table_buffer: list[str] = []

    def flush_table():
        nonlocal table_buffer
        if table_buffer:
            add_markdown_table(doc, table_buffer)
            table_buffer = []

    for raw in lines:
        line = raw.rstrip()

        if not line.strip():
            flush_table()
            continue

        if line.strip().startswith("<!--"):
            continue

        if line.strip().startswith("|") and "|" in line.strip()[1:]:
            table_buffer.append(line)
            continue

        flush_table()

        if line.startswith("# "):
            p = doc.add_paragraph(normalize_inline_markdown(line[2:].strip()), style="Heading 1")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith("## "):
            p = doc.add_paragraph(normalize_inline_markdown(line[3:].strip()), style="Heading 2")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith("### "):
            p = doc.add_paragraph(normalize_inline_markdown(line[4:].strip()), style="Heading 3")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif re.match(r"^\s*[-*]\s+", line):
            add_plain_paragraph(doc, re.sub(r"^\s*[-*]\s+", "", line))
        else:
            add_plain_paragraph(doc, line)

    flush_table()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--md", required=True)
    parser.add_argument("--out", required=True)
    args = parser.parse_args()

    md_path = Path(args.md)
    out_path = Path(args.out)

    if not md_path.exists():
        raise SystemExit(f"[ERROR] Missing MD: {md_path}")

    md_to_docx(md_path, out_path)
    print("[INFO] DOCX:", out_path)


if __name__ == "__main__":
    main()
