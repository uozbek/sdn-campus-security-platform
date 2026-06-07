#!/usr/bin/env python3
from __future__ import annotations

import argparse
import csv
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


TABLE_MAP = {
    "Tablo 1": "tables/table_canonical_runtime_validation_summary.csv",
    "Tablo 2": "tables/table_controller_action_distribution.csv",
    "Tablo 3": "tables/table_final_top20_prediction_distribution.csv",
    "Tablo 4": "tables/table_protocol_aware_final_policy_distribution.csv",
    "Tablo 5": "tables/table_enforcement_action_summary.csv",
    "Tablo 6": "tables/table_flow_level_model_controller_comparison.csv",
    "Tablo 7": "tables/table_protocol_aware_attack_probability_summary.csv",
}

FIGURE_MAP = {
    "Şekil 1": "figures/fig_controller_action_distribution.png",
    "Şekil 2": "figures/fig_final_top20_prediction_distribution.png",
    "Şekil 3": "figures/fig_protocol_aware_final_policy_distribution.png",
    "Şekil 4": "figures/fig_enforcement_action_summary.png",
}


def set_cell_shading(cell, fill: str = "D9EAF7") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(8)


def add_csv_table(doc: Document, csv_path: Path, max_rows: int | None = None) -> None:
    if not csv_path.exists():
        p = doc.add_paragraph()
        p.add_run(f"[UYARI: Tablo dosyası bulunamadı: {csv_path}]").bold = True
        return

    with csv_path.open("r", encoding="utf-8") as f:
        rows = list(csv.reader(f))

    if not rows:
        doc.add_paragraph(f"[UYARI: Tablo boş: {csv_path}]")
        return

    if max_rows is not None and len(rows) > max_rows + 1:
        rows = rows[: max_rows + 1]
        truncated = True
    else:
        truncated = False

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_text(cell, value, bold=(i == 0))
            if i == 0:
                set_cell_shading(cell)
        if i == 0:
            set_repeat_table_header(table.rows[0])

    if truncated:
        doc.add_paragraph(
            "Not: Tablo Word düzeni için kısaltılmıştır. Tam tablo CSV dosyasında yer almaktadır."
        ).italic = True


def add_caption(doc: Document, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.bold = True
    run.font.size = Pt(10)


def add_figure(doc: Document, image_path: Path, caption: str) -> None:
    if not image_path.exists():
        p = doc.add_paragraph()
        p.add_run(f"[UYARI: Şekil dosyası bulunamadı: {image_path}]").bold = True
        return

    doc.add_picture(str(image_path), width=Inches(6.2))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, caption)


def clean_inline_markdown(text: str) -> str:
    text = text.replace("`", "")
    text = text.replace("**", "")
    return text


def add_markdown_paragraph(doc: Document, line: str) -> None:
    stripped = line.strip()

    if not stripped:
        return

    if stripped.startswith("# "):
        doc.add_heading(clean_inline_markdown(stripped[2:]), level=1)
    elif stripped.startswith("## "):
        doc.add_heading(clean_inline_markdown(stripped[3:]), level=2)
    elif stripped.startswith("### "):
        doc.add_heading(clean_inline_markdown(stripped[4:]), level=3)
    elif stripped.startswith("- "):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(clean_inline_markdown(stripped[2:]))
    elif re.match(r"^\d+\.\s+", stripped):
        p = doc.add_paragraph(style="List Number")
        p.add_run(clean_inline_markdown(re.sub(r"^\d+\.\s+", "", stripped)))
    elif stripped.startswith("> "):
        p = doc.add_paragraph()
        run = p.add_run(clean_inline_markdown(stripped[2:]))
        run.italic = True
    elif stripped == "---":
        doc.add_page_break()
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(clean_inline_markdown(stripped))


def configure_styles(doc: Document) -> None:
    styles = doc.styles

    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.bold = True

    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 3"].font.size = Pt(12)


def build_docx(markdown_path: Path, artifact_dir: Path, output_path: Path) -> None:
    doc = Document()
    configure_styles(doc)

    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    lines = markdown_path.read_text(encoding="utf-8").splitlines()

    skip_next_caption = False

    for current_index, line in enumerate(lines):
        stripped = line.strip()

        table_match = re.search(r"$begin:math:display$TABLO EKLEME NOKTASI — \(Tablo \\d\+\)$end:math:display$", stripped)
        figure_match = re.search(r"$begin:math:display$ŞEKİL EKLEME NOKTASI — \(Şekil \\d\+\)$end:math:display$", stripped)

        if table_match:
            table_no = table_match.group(1)
            csv_rel = TABLE_MAP.get(table_no)
            if csv_rel:
                doc.add_paragraph()
                add_caption(doc, f"{table_no}. {extract_title_after_marker(lines, current_index)}")
                max_rows = 12 if table_no in {"Tablo 1", "Tablo 6", "Tablo 7"} else None
                add_csv_table(doc, artifact_dir / csv_rel, max_rows=max_rows)
                doc.add_paragraph()
            continue

        if figure_match:
            fig_no = figure_match.group(1)
            img_rel = FIGURE_MAP.get(fig_no)
            if img_rel:
                doc.add_paragraph()
                add_figure(doc, artifact_dir / img_rel, f"{fig_no}. {extract_title_after_marker(lines, current_index)}")
                doc.add_paragraph()
            continue

        # Avoid duplicating the textual "Dosya:" placeholder lines.
        if stripped.startswith("> Dosya:") or stripped.startswith("> Başlık:"):
            continue

        add_markdown_paragraph(doc, line)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def extract_title_after_marker(lines: list[str], current_index: int) -> str:
    # The title is available on a following placeholder line as "> Başlık: ..."
    for follow in lines[current_index: current_index + 8]:
        s = follow.strip()
        if "Başlık:" in s:
            return s.split("Başlık:", 1)[1].strip().replace("`", "").replace("**", "")
    return "Deneysel çıktı"


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--markdown", required=True)
    parser.add_argument("--artifact-dir", required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()

    build_docx(
        markdown_path=Path(args.markdown),
        artifact_dir=Path(args.artifact_dir),
        output_path=Path(args.output),
    )

    print(f"[INFO] Written: {args.output}")


if __name__ == "__main__":
    main()
