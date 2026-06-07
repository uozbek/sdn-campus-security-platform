#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
from datetime import datetime
from pathlib import Path

from docx import Document


def cm(x) -> float:
    return round(float(x.cm), 2)


def status_close(actual: float, expected: float, tolerance: float) -> str:
    return "ok" if abs(actual - expected) <= tolerance else "check"


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--docx",
        default="docs/tez_ana_taslak_tr_frontmatter_ozetli.docx",
    )
    parser.add_argument(
        "--rules-json",
        default="docs/sau_fbe_format_rules.json",
    )
    parser.add_argument(
        "--output-md",
        default="docs/sau_fbe_format_check_report.md",
    )
    parser.add_argument(
        "--output-json",
        default="docs/sau_fbe_format_check_report.json",
    )
    args = parser.parse_args()

    docx_path = Path(args.docx)
    rules_path = Path(args.rules_json)

    if not docx_path.exists():
        raise SystemExit(f"[ERROR] DOCX not found: {docx_path}")
    if not rules_path.exists():
        raise SystemExit(f"[ERROR] Rules JSON not found: {rules_path}")

    rules = json.loads(rules_path.read_text(encoding="utf-8"))
    doc = Document(docx_path)
    section = doc.sections[0]

    tolerance = float(rules["margins"].get("tolerance_cm", 0.15))

    actual = {
        "page": {
            "width_cm": cm(section.page_width),
            "height_cm": cm(section.page_height),
        },
        "margins": {
            "top_cm": cm(section.top_margin),
            "bottom_cm": cm(section.bottom_margin),
            "left_cm": cm(section.left_margin),
            "right_cm": cm(section.right_margin),
            "header_distance_cm": cm(section.header_distance),
            "footer_distance_cm": cm(section.footer_distance),
        },
        "counts": {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "inline_shapes": len(doc.inline_shapes),
            "sections": len(doc.sections),
        },
    }

    checks = []

    for key in ["width_cm", "height_cm"]:
        exp = float(rules["page"][key])
        act = float(actual["page"][key])
        checks.append({
            "control": f"Sayfa {key}",
            "expected": exp,
            "actual": act,
            "status": status_close(act, exp, 0.05),
        })

    for key in ["top_cm", "bottom_cm", "left_cm", "right_cm", "header_distance_cm", "footer_distance_cm"]:
        exp = float(rules["margins"][key])
        act = float(actual["margins"][key])
        checks.append({
            "control": f"Kenar boşluğu {key}",
            "expected": exp,
            "actual": act,
            "status": status_close(act, exp, tolerance),
        })

    # Normal style
    try:
        normal = doc.styles["Normal"]
        act_font = normal.font.name or "not-set"
        act_size = normal.font.size.pt if normal.font.size else None
    except Exception:
        act_font = "error"
        act_size = None

    exp_normal = rules["styles"].get("Normal", {})
    checks.append({
        "control": "Normal stil yazı tipi",
        "expected": exp_normal.get("font_name"),
        "actual": act_font,
        "status": "ok" if (not exp_normal.get("font_name") or exp_normal.get("font_name") == act_font) else "check",
    })
    checks.append({
        "control": "Normal stil punto",
        "expected": exp_normal.get("font_size_pt"),
        "actual": act_size,
        "status": "ok" if (exp_normal.get("font_size_pt") in [None, act_size]) else "check",
    })

    # Required text presence
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    def norm(value: str) -> str:
        value = (value or "").strip().lower()
        table = str.maketrans({
            "İ": "i", "I": "i", "ı": "i",
            "Ş": "s", "ş": "s",
            "Ğ": "g", "ğ": "g",
            "Ü": "u", "ü": "u",
            "Ö": "o", "ö": "o",
            "Ç": "c", "ç": "c",
        })
        return value.translate(table)

    texts_norm = [norm(t) for t in texts]

    required_presence = [
        "Özet",
        "Abstract",
        "İçindekiler",
        "Tablo Listesi",
        "Şekil Listesi",
        "Kaynakça",
        "Bölüm 1. Giriş",
        "Bölüm 2. Kavramsal ve Kuramsal Arka Plan",
        "Bölüm 3. Literatür Taraması ve İlgili Çalışmalar",
        "Bölüm 4. Yöntem ve Çalışma Zamanı Doğrulama",
        "Bölüm 5. Tartışma, Sınırlılıklar ve Gelecek Çalışmalar",
        "Bölüm 6. Sonuç ve Öneriler",
    ]

    aliases = {
        "Özet": ["özet"],
        "Abstract": ["abstract", "summary"],
        "İçindekiler": ["içindekiler", "icindekiler"],
        "Tablo Listesi": ["tablo listesi"],
        "Şekil Listesi": ["şekil listesi", "sekil listesi"],
        "Kaynakça": ["kaynakça", "kaynaklar"],
    }

    for item in required_presence:
        search_terms = aliases.get(item, [item])
        search_terms = [norm(term) for term in search_terms]
        found = any(any(term in t for term in search_terms) for t in texts_norm)
        checks.append({
            "control": f"Metin/Başlık varlığı: {item}",
            "expected": "present",
            "actual": "present" if found else "missing",
            "status": "ok" if found else "check",
        })

    result = {
        "generated_at_utc": datetime.utcnow().isoformat(),
        "docx": str(docx_path),
        "rules_json": str(rules_path),
        "actual": actual,
        "checks": checks,
        "summary": {
            "ok": sum(1 for c in checks if c["status"] == "ok"),
            "check": sum(1 for c in checks if c["status"] != "ok"),
            "total": len(checks),
        },
    }

    out_json = Path(args.output_json)
    out_json.parent.mkdir(parents=True, exist_ok=True)
    out_json.write_text(json.dumps(result, indent=2, ensure_ascii=False), encoding="utf-8")

    md = []
    md.append("# SAÜ FBE Format Kontrol Raporu")
    md.append("")
    md.append(f"- Generated at UTC: `{result['generated_at_utc']}`")
    md.append(f"- DOCX: `{docx_path}`")
    md.append(f"- Rules JSON: `{rules_path}`")
    md.append("")
    md.append("## 1. Özet")
    md.append("")
    md.append(f"- OK: `{result['summary']['ok']}`")
    md.append(f"- Check: `{result['summary']['check']}`")
    md.append(f"- Total: `{result['summary']['total']}`")
    md.append("")
    md.append("## 2. Kontroller")
    md.append("")
    md.append("| Kontrol | Beklenen | Mevcut | Durum |")
    md.append("|---|---:|---:|---|")
    for c in checks:
        md.append(f"| {c['control']} | {c['expected']} | {c['actual']} | {c['status']} |")

    md.append("")
    md.append("## 3. Not")
    md.append("")
    md.append("Bu rapor, SAÜ FBE tez şablonundan çıkarılan DOCX ayarları ile mevcut ana tez taslağını karşılaştırır. Tablo/şekil başlığı konumu, kaynakça biçimi ve ön kapak/onay sayfaları için daha ayrıntılı denetimler sonraki aşamada eklenecektir.")

    out_md = Path(args.output_md)
    out_md.parent.mkdir(parents=True, exist_ok=True)
    out_md.write_text("\n".join(md), encoding="utf-8")

    print("[INFO] MD:", out_md)
    print("[INFO] JSON:", out_json)
    print("[INFO] Summary:", result["summary"])


if __name__ == "__main__":
    main()
