#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document


def clean(value) -> str:
    if value is None:
        return ""
    s = str(value)
    if s.lower() == "nan":
        return ""
    return s.strip()


def normalize_text(s: str) -> str:
    s = clean(s).lower()
    tr_map = str.maketrans({
        "ç": "c", "ğ": "g", "ı": "i", "ö": "o", "ş": "s", "ü": "u",
        "Ç": "c", "Ğ": "g", "İ": "i", "I": "i", "Ö": "o", "Ş": "s", "Ü": "u",
    })
    s = s.translate(tr_map)
    s = re.sub(r"[^a-z0-9]+", " ", s)
    return re.sub(r"\s+", " ", s).strip()


def extract_surname_from_authors(authors: str) -> str:
    authors = clean(authors)
    if not authors:
        return ""

    # BibTeX: "Surname, Name and Surname2, Name2"
    first = re.split(r"\s+and\s+|;", authors)[0].strip()

    if "," in first:
        return first.split(",", 1)[0].strip()

    tokens = first.split()
    return tokens[-1].strip() if tokens else ""


def extract_docx_text(docx_path: Path) -> tuple[str, list[str]]:
    doc = Document(docx_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # Tablo hücrelerini de dahil et.
    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    table_texts.append(t)

    all_parts = paragraphs + table_texts
    return "\n".join(all_parts), all_parts


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--references-csv", default="docs/references_zotero_apa_like_reviewed.csv")
    parser.add_argument("--thesis-docx", default="docs/tez_ana_taslak_tr_guncel_sau_fbe.docx")
    parser.add_argument("--output-md", default="docs/bibliography_reference_usage_audit_zotero_reviewed.md")
    parser.add_argument("--output-csv", default="docs/bibliography_reference_usage_audit_zotero_reviewed.csv")
    parser.add_argument("--output-json", default="docs/bibliography_reference_usage_audit_zotero_reviewed.json")
    args = parser.parse_args()

    refs_path = Path(args.references_csv)
    docx_path = Path(args.thesis_docx)

    if not refs_path.exists():
        raise SystemExit(f"[ERROR] references CSV not found: {refs_path}")
    if not docx_path.exists():
        raise SystemExit(f"[ERROR] thesis DOCX not found: {docx_path}")

    refs = pd.read_csv(refs_path).fillna("")
    full_text, parts = extract_docx_text(docx_path)
    norm_full = normalize_text(full_text)

    rows = []

    for _, r in refs.iterrows():
        key = clean(r.get("zotero_key") or r.get("id"))
        authors = clean(r.get("authors"))
        formatted_authors = clean(r.get("formatted_authors"))
        title = clean(r.get("title"))
        year = clean(r.get("year"))
        doi = clean(r.get("doi"))
        url = clean(r.get("url"))

        surname = extract_surname_from_authors(authors) or extract_surname_from_authors(formatted_authors)
        year4 = re.search(r"(19|20)\d{2}", year)
        year4 = year4.group(0) if year4 else ""

        norm_surname = normalize_text(surname)
        norm_title = normalize_text(title)
        title_words = [w for w in norm_title.split() if len(w) >= 5]
        title_probe = " ".join(title_words[:5])

        surname_hit = bool(norm_surname and norm_surname in norm_full)
        year_hit = bool(year4 and year4 in full_text)
        title_hit = bool(title_probe and title_probe in norm_full)

        # APA inline olasılığı: soyad + yıl aynı dokümanda geçiyor mu?
        inline_like_hit = surname_hit and year_hit

        # DOI veya URL kaynakça bölümünde geçiyor olabilir.
        doi_hit = bool(doi and doi.lower() in full_text.lower())
        url_hit = bool(url and url.lower() in full_text.lower())

        if inline_like_hit:
            status = "likely_cited_or_mentioned"
        elif title_hit:
            status = "title_mentioned"
        elif doi_hit or url_hit:
            status = "bibliography_only_or_url_hit"
        else:
            status = "not_cited_or_not_mentioned"

        rows.append({
            "zotero_key": key,
            "year": year4,
            "surname": surname,
            "title": title,
            "final_decision": clean(r.get("final_decision")),
            "relevance_to_this_thesis": clean(r.get("relevance_to_this_thesis")),
            "surname_hit": surname_hit,
            "year_hit": year_hit,
            "inline_like_hit": inline_like_hit,
            "title_hit": title_hit,
            "doi_hit": doi_hit,
            "url_hit": url_hit,
            "status": status,
        })

    out = pd.DataFrame(rows)
    Path(args.output_csv).parent.mkdir(parents=True, exist_ok=True)
    out.to_csv(args.output_csv, index=False)

    summary = {
        "generated_at_utc": datetime.utcnow().isoformat(),
        "references_csv": str(refs_path),
        "thesis_docx": str(docx_path),
        "reference_count": int(len(out)),
        "status_counts": out["status"].value_counts().to_dict(),
    }

    Path(args.output_json).write_text(json.dumps(summary, indent=2, ensure_ascii=False), encoding="utf-8")

    md = []
    md.append("# Bibliography Reference Usage Audit — Zotero Reviewed")
    md.append("")
    md.append(f"- Generated at UTC: `{summary['generated_at_utc']}`")
    md.append(f"- References CSV: `{refs_path}`")
    md.append(f"- Thesis DOCX: `{docx_path}`")
    md.append(f"- Reference count: `{len(out)}`")
    md.append("")
    md.append("## 1. Status Counts")
    md.append("")
    md.append("| Status | Count |")
    md.append("|---|---:|")
    for k, v in out["status"].value_counts().items():
        md.append(f"| {k} | {v} |")

    md.append("")
    md.append("## 2. Not Cited / Not Mentioned")
    md.append("")
    missing = out[out["status"].eq("not_cited_or_not_mentioned")]
    if missing.empty:
        md.append("No uncited/unmentioned references detected.")
    else:
        md.append("| Key | Year | Surname | Decision | Title |")
        md.append("|---|---:|---|---|---|")
        for _, r in missing.iterrows():
            title = clean(r["title"]).replace("|", "\\|")
            md.append(f"| {r['zotero_key']} | {r['year']} | {r['surname']} | {r['final_decision']} | {title[:180]} |")

    md.append("")
    md.append("## 3. Likely Cited / Mentioned")
    md.append("")
    cited = out[out["status"].ne("not_cited_or_not_mentioned")]
    md.append("| Key | Status | Year | Surname | Title |")
    md.append("|---|---|---:|---|---|")
    for _, r in cited.iterrows():
        title = clean(r["title"]).replace("|", "\\|")
        md.append(f"| {r['zotero_key']} | {r['status']} | {r['year']} | {r['surname']} | {title[:160]} |")

    Path(args.output_md).write_text("\n".join(md), encoding="utf-8")

    print("[INFO] MD:", args.output_md)
    print("[INFO] CSV:", args.output_csv)
    print("[INFO] JSON:", args.output_json)
    print("[INFO] Status counts:", summary["status_counts"])


if __name__ == "__main__":
    main()
