#!/usr/bin/env python3
from __future__ import annotations

import argparse
import copy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Times New Roman"
        styles[style_name].font.bold = True

    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 3"].font.size = Pt(12)


def add_centered(doc: Document, text: str, size: int = 12, bold: bool = False, space_after: int = 6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_justified(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    return p


def add_frontmatter(doc: Document) -> None:
    # Cover placeholder
    add_centered(doc, "T.C.", 12, True)
    add_centered(doc, "[ÜNİVERSİTE ADI]", 12, True)
    add_centered(doc, "[ENSTİTÜ ADI]", 12, True)
    add_centered(doc, "[ANABİLİM DALI]", 12, True)

    for _ in range(4):
        doc.add_paragraph()

    add_centered(
        doc,
        "YAZILIM TANIMLI KAMPÜS AĞLARINDA MAKİNE ÖĞRENMESİ DESTEKLİ DDOS TESPİTİ VE ÖNLEME",
        14,
        True,
        12,
    )

    for _ in range(3):
        doc.add_paragraph()

    add_centered(doc, "Doktora Tezi", 12, True)
    add_centered(doc, "[Ad Soyad]", 12, True)

    for _ in range(3):
        doc.add_paragraph()

    add_centered(doc, "Danışman", 12, True)
    add_centered(doc, "[Danışman Unvanı Ad Soyad]", 12, True)

    for _ in range(5):
        doc.add_paragraph()

    add_centered(doc, "[Şehir]", 12, False)
    add_centered(doc, "2026", 12, False)
    doc.add_page_break()

    # Declaration placeholders
    doc.add_heading("Tez Onay Sayfası", level=1)
    add_justified(doc, "Bu sayfa, enstitü tez yazım kılavuzuna uygun şekilde daha sonra doldurulacaktır.")
    doc.add_page_break()

    doc.add_heading("Beyan", level=1)
    add_justified(doc, "Bu tez çalışmasının bilimsel etik kurallara uygun olarak hazırlandığını beyan ederim. Bu bölüm, ilgili enstitü formatına göre düzenlenecektir.")
    doc.add_page_break()

    doc.add_heading("Özet", level=1)
    add_justified(doc, "Bu bölümde tez çalışmasının amacı, yöntemi, deneysel bulguları ve sonuçları Türkçe olarak özetlenecektir. Mevcut çalışma, SDN tabanlı kampüs ağı benzeri bir ortamda makine öğrenmesi destekli DDoS tespit ve önleme prototipinin geliştirilmesine odaklanmaktadır.")
    add_justified(doc, "Anahtar Kelimeler: Yazılım Tanımlı Ağlar, DDoS, Makine Öğrenmesi, IDS, IPS, Ryu, OpenFlow")
    doc.add_page_break()

    doc.add_heading("Abstract", level=1)
    add_justified(doc, "This section will summarize the aim, methodology, experimental findings, and conclusions of the dissertation in English. The study focuses on developing a machine-learning-assisted DDoS detection and prevention prototype in an SDN-based campus-like network environment.")
    add_justified(doc, "Keywords: Software-Defined Networking, DDoS, Machine Learning, IDS, IPS, Ryu, OpenFlow")
    doc.add_page_break()

    doc.add_heading("İçindekiler", level=1)
    add_justified(doc, "Not: İçindekiler tablosu Word üzerinde References > Table of Contents menüsünden otomatik olarak güncellenmelidir.")
    add_justified(doc, "Bölüm 1. Giriş")
    add_justified(doc, "Bölüm 2. Kavramsal ve Kuramsal Arka Plan")
    add_justified(doc, "Bölüm 3. Literatür Taraması ve İlgili Çalışmalar")
    add_justified(doc, "Bölüm 4. Yöntem ve Çalışma Zamanı Doğrulama")
    add_justified(doc, "Bölüm 5. Tartışma, Sınırlılıklar ve Gelecek Çalışmalar")
    add_justified(doc, "Bölüm 6. Sonuç ve Öneriler")
    doc.add_page_break()

    doc.add_heading("Tablo Listesi", level=1)
    add_justified(doc, "Tablo 1. Canonical tekrarlı çalışma zamanı doğrulama sonuçları")
    add_justified(doc, "Tablo 2. Denetleyici taraflı IDS/IPS politika aksiyon dağılımı")
    add_justified(doc, "Tablo 3. Final XGBoost Top-20 modelinin runtime tahmin dağılımı")
    add_justified(doc, "Tablo 4. Protocol-aware final policy dağılımı")
    add_justified(doc, "Tablo 5. SDN denetleyicisi önleme aksiyon özeti")
    add_justified(doc, "Tablo 6. Flow-level model-controller karşılaştırması")
    add_justified(doc, "Tablo 7. Flow-level saldırı olasılığı ve final action özeti")
    doc.add_page_break()

    doc.add_heading("Şekil Listesi", level=1)
    add_justified(doc, "Şekil 1. Denetleyici taraflı politika aksiyon dağılımı")
    add_justified(doc, "Şekil 2. Runtime BENIGN/ATTACK tahmin dağılımı")
    add_justified(doc, "Şekil 3. Protocol-aware final policy aksiyon dağılımı")
    add_justified(doc, "Şekil 4. SDN denetleyicisi önleme aksiyon özeti")
    doc.add_page_break()


def append_source_doc(target: Document, source_path: Path) -> None:
    source = Document(source_path)
    target_body = target.element.body

    for child in source.element.body:
        if child.tag.endswith("}sectPr"):
            continue
        target_body.append(copy.deepcopy(child))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True, help="Existing main thesis draft docx")
    parser.add_argument("--output", required=True, help="Output docx with front matter")
    args = parser.parse_args()

    input_path = Path(args.input)
    output_path = Path(args.output)

    if not input_path.exists():
        raise SystemExit(f"[ERROR] Input not found: {input_path}")

    doc = Document()
    configure_doc(doc)

    # Remove default empty paragraph.
    if len(doc.paragraphs) == 1 and not doc.paragraphs[0].text.strip():
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)

    add_frontmatter(doc)
    append_source_doc(doc, input_path)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    print("[INFO] Written:", output_path)


if __name__ == "__main__":
    main()
