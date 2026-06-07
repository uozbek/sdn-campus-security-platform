#!/usr/bin/env python3
from __future__ import annotations

import argparse
from pathlib import Path

import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


def add_centered(doc, text, size=12, bold=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_left_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_page_setup(doc):
    for s in doc.sections:
        s.page_width = Cm(21.0)
        s.page_height = Cm(29.7)
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(4.0)
        s.right_margin = Cm(2.5)
        s.header_distance = Cm(1.25)
        s.footer_distance = Cm(1.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)


def add_cover(doc, meta, inner=False):
    add_centered(doc, meta["university"], size=14, bold=True, space_after=2)
    add_centered(doc, meta["institute"], size=14, bold=True, space_after=36)

    add_centered(doc, meta["title_tr"], size=14, bold=True, space_after=36)

    add_centered(doc, meta["author"], size=12, bold=True, space_after=24)
    add_centered(doc, meta["degree_type_tr"], size=12, bold=True, space_after=36)

    add_centered(doc, meta.get("department", ""), size=12, bold=False, space_after=8)
    add_centered(doc, meta.get("program", ""), size=12, bold=False, space_after=24)

    if inner:
        add_centered(doc, f"Danışman: {meta.get('advisor', 'TODO_DANISMAN')}", size=12, bold=False, space_after=6)
        if meta.get("co_advisor"):
            add_centered(doc, f"İkinci Danışman: {meta.get('co_advisor')}", size=12, bold=False, space_after=6)

    doc.add_paragraph("")
    doc.add_paragraph("")
    add_centered(doc, f"{meta.get('city', 'SAKARYA')} - {meta.get('year', '2026')}", size=12, bold=False, space_after=0)


def add_approval_page(doc, meta):
    add_left_heading(doc, "ONAY SAYFASI", level=1)
    doc.add_paragraph(
        f"{meta.get('author', 'TODO_AD_SOYAD')} tarafından hazırlanan "
        f"“{meta.get('title_tr', '')}” başlıklı bu çalışma, ilgili jüri tarafından "
        f"{meta.get('degree_type_tr', 'DOKTORA TEZİ')} olarak kabul edilmiştir."
    )
    doc.add_paragraph("")
    doc.add_paragraph("Jüri Üyeleri:")
    for member in meta.get("jury", []):
        doc.add_paragraph(member, style="List Bullet")
    doc.add_paragraph("")
    doc.add_paragraph("Bu sayfa, nihai teslim öncesinde enstitü tarafından istenen imza/onay formatına göre düzenlenmelidir.")


def add_declaration(doc):
    add_left_heading(doc, "BEYAN", level=1)
    doc.add_paragraph(
        "Bu tez çalışmasının kendi çalışmam olduğunu, tezin planlanmasından yazımına kadar bütün aşamalarda "
        "bilimsel etik ilke ve kurallara uygun davrandığımı, bu tezdeki bütün bilgileri akademik ve etik kurallar "
        "içinde elde ettiğimi, kullanılan tüm kaynaklara atıf yaptığımı beyan ederim."
    )
    doc.add_paragraph("")
    doc.add_paragraph("İmza")
    doc.add_paragraph("Ad Soyad")


def add_acknowledgement(doc):
    add_left_heading(doc, "TEŞEKKÜR", level=1)
    doc.add_paragraph(
        "Bu bölümde tez çalışmasının yürütülmesi sürecinde katkı sağlayan danışman, jüri üyeleri, kurumlar ve "
        "destek veren kişiler için teşekkür metni nihai aşamada düzenlenecektir."
    )


def add_lists_placeholders(doc):
    for title, desc in [
        ("İÇİNDEKİLER", "Bu bölüm Microsoft Word üzerinde otomatik içindekiler alanı ile güncellenmelidir."),
        ("KISALTMALAR", "Bu bölümde tezde kullanılan kısaltmalar listelenecektir."),
        ("SİMGELER", "Bu bölümde tezde kullanılan simgeler listelenecektir."),
        ("TABLO LİSTESİ", "Bu bölüm Microsoft Word üzerinde otomatik tablo listesi alanı ile güncellenmelidir."),
        ("ŞEKİL LİSTESİ", "Bu bölüm Microsoft Word üzerinde otomatik şekil listesi alanı ile güncellenmelidir."),
    ]:
        add_left_heading(doc, title, level=1)
        doc.add_paragraph(desc)


def add_abstract_pages(doc, meta):
    add_left_heading(doc, "ÖZET", level=1)
    doc.add_paragraph(
        "Bu tez çalışması, yazılım tanımlı kampüs ağlarında DDoS saldırılarının makine öğrenmesi destekli olarak "
        "tespit edilmesi ve SDN denetleyicisi üzerinden aktif önleme aksiyonlarına dönüştürülmesi problemine odaklanmaktadır. "
        "Geliştirilen prototipte Mininet, Open vSwitch, Ryu denetleyicisi, FastAPI tabanlı model servisi ve OpenFlow "
        "kuralları birlikte kullanılmıştır."
    )
    doc.add_paragraph("Anahtar kelimeler: " + ", ".join(meta.get("keywords_tr", [])))

    add_left_heading(doc, "ABSTRACT", level=1)
    doc.add_paragraph(
        "This thesis focuses on machine-learning-assisted detection and SDN-controller-based mitigation of DDoS attacks "
        "in software-defined campus networks. The proposed prototype integrates Mininet, Open vSwitch, the Ryu controller, "
        "a FastAPI-based inference service, and OpenFlow-based enforcement actions."
    )
    doc.add_paragraph("Keywords: " + ", ".join(meta.get("keywords_en", [])))


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--metadata", default="docs/thesis_metadata.yaml")
    parser.add_argument("--output", default="docs/sau_fbe_frontmatter_generated.docx")
    args = parser.parse_args()

    meta_path = Path(args.metadata)
    if not meta_path.exists():
        raise SystemExit(f"[ERROR] Metadata file not found: {meta_path}")

    meta = yaml.safe_load(meta_path.read_text(encoding="utf-8"))

    doc = Document()
    add_page_setup(doc)

    add_cover(doc, meta, inner=False)
    doc.add_page_break()
    add_cover(doc, meta, inner=True)
    doc.add_page_break()
    add_approval_page(doc, meta)
    doc.add_page_break()
    add_declaration(doc)
    doc.add_page_break()
    add_acknowledgement(doc)
    doc.add_page_break()
    add_lists_placeholders(doc)
    doc.add_page_break()
    add_abstract_pages(doc, meta)

    out = Path(args.output)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)

    print("[INFO] Written:", out)


if __name__ == "__main__":
    main()
