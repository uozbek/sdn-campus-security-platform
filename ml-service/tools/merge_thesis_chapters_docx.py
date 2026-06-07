#!/usr/bin/env python3
from __future__ import annotations

import argparse
import copy
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Times New Roman"
        styles[style_name].font.bold = True

    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 3"].font.size = Pt(12)


def append_doc_body(target: Document, source_path: Path, add_page_break: bool = True) -> None:
    source = Document(source_path)

    if add_page_break and len(target.element.body) > 0:
        target.add_page_break()

    target_body = target.element.body

    # Remove default empty paragraph only if target is still effectively empty.
    if len(target.paragraphs) == 1 and not target.paragraphs[0].text.strip() and len(target.tables) == 0:
        p = target.paragraphs[0]._element
        p.getparent().remove(p)

    for child in source.element.body:
        # Skip section properties from source documents.
        if child.tag.endswith("}sectPr"):
            continue
        target_body.append(copy.deepcopy(child))


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Merge thesis chapter DOCX files into one main thesis draft."
    )
    parser.add_argument("--output", required=True)
    parser.add_argument("--chapter", action="append", required=True, help="Chapter DOCX path. Repeatable.")
    args = parser.parse_args()

    output = Path(args.output)
    chapters = [Path(x) for x in args.chapter]

    missing = [str(p) for p in chapters if not p.exists()]
    if missing:
        raise SystemExit("[ERROR] Missing chapter files:\n" + "\n".join(missing))

    merged = Document()
    configure_doc(merged)

    for i, ch in enumerate(chapters):
        append_doc_body(merged, ch, add_page_break=(i > 0))

    output.parent.mkdir(parents=True, exist_ok=True)
    merged.save(output)

    print("[INFO] Written:", output)
    print("[INFO] Chapters merged:", len(chapters))
    for ch in chapters:
        print(" -", ch)


if __name__ == "__main__":
    main()
