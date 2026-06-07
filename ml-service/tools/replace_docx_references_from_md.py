#!/usr/bin/env python3
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.shared import Pt


def clean_line(line: str) -> str:
    return line.strip()


def load_references(md_path: Path) -> list[str]:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    refs = []
    for line in lines:
        line = clean_line(line)
        if line.startswith("- "):
            refs.append(line[2:].strip())
    return refs


def is_heading_text(text: str, heading: str) -> bool:
    return text.strip().lower() == heading.strip().lower()


def clear_paragraph(paragraph):
    for run in paragraph.runs:
        run.text = ""


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input-docx", required=True)
    parser.add_argument("--references-md", default="docs/references_apa_like.md")
    parser.add_argument("--output-docx", required=True)
    args = parser.parse_args()

    input_docx = Path(args.input_docx)
    refs_md = Path(args.references_md)
    output_docx = Path(args.output_docx)

    if not input_docx.exists():
        raise SystemExit(f"[ERROR] Input DOCX not found: {input_docx}")
    if not refs_md.exists():
        raise SystemExit(f"[ERROR] References MD not found: {refs_md}")

    refs = load_references(refs_md)
    if not refs:
        raise SystemExit(f"[ERROR] No reference lines found in: {refs_md}")

    doc = Document(input_docx)

    paragraphs = doc.paragraphs
    ref_idx = None
    for i, p in enumerate(paragraphs):
        if is_heading_text(p.text, "Kaynakça") or is_heading_text(p.text, "Kaynaklar"):
            ref_idx = i
            break

    if ref_idx is None:
        # If no existing heading found, append at end.
        doc.add_page_break()
        heading = doc.add_paragraph("Kaynakça")
        try:
            heading.style = "Heading 1"
        except Exception:
            pass
        ref_idx = len(doc.paragraphs) - 1
    else:
        # Normalize heading text.
        clear_paragraph(paragraphs[ref_idx])
        run = paragraphs[ref_idx].add_run("Kaynakça")
        run.bold = True

    # Remove existing content after Kaynakça until Ekler/Özgeçmiş or EOF.
    stop_headings = {"ekler", "özgeçmiş", "ozgecmis"}
    start = ref_idx + 1
    stop = len(doc.paragraphs)

    for i in range(start, len(doc.paragraphs)):
        t = doc.paragraphs[i].text.strip().lower()
        if t in stop_headings:
            stop = i
            break

    # Clear old paragraphs in range instead of physically deleting to avoid XML complexity.
    for i in range(start, stop):
        clear_paragraph(doc.paragraphs[i])

    insert_after = ref_idx

    # Add a note paragraph, then references.
    note = doc.paragraphs[insert_after].insert_paragraph_before("")
    # Because python-docx lacks insert_after, we append references at end if needed.
    # Simpler reliable approach: append clean Kaynakça section at document end if old section cannot be physically removed.
    # To avoid duplicated visible content, old content was cleared above.

    for ref in refs:
        p = doc.add_paragraph(ref)
        try:
            p.style = doc.styles["Normal"]
        except Exception:
            pass
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.left_indent = Pt(0)
        p.paragraph_format.space_after = Pt(6)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)

    print("[INFO] Input :", input_docx)
    print("[INFO] Refs  :", refs_md)
    print("[INFO] Output:", output_docx)
    print("[INFO] Reference count:", len(refs))


if __name__ == "__main__":
    main()
