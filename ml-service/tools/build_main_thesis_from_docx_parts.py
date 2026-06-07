#!/usr/bin/env python3
from __future__ import annotations

import argparse
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt


def apply_sau_fbe_page_setup(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(4.0)
        section.right_margin = Cm(2.5)
        section.header_distance = Cm(1.25)
        section.footer_distance = Cm(1.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)


def append_docx(master: Document, part_path: Path, page_break: bool = True) -> None:
    part = Document(part_path)

    if page_break and len(master.paragraphs) > 0:
        master.add_page_break()

    body = master.element.body
    for element in part.element.body:
        if element.tag.endswith("sectPr"):
            continue
        body.append(deepcopy(element))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--frontmatter", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--part", action="append", required=True, help="DOCX part path. Repeatable.")
    args = parser.parse_args()

    frontmatter = Path(args.frontmatter)
    output = Path(args.output)

    if not frontmatter.exists():
        raise SystemExit(f"[ERROR] Frontmatter not found: {frontmatter}")

    doc = Document(frontmatter)

    for part in args.part:
        p = Path(part)
        if not p.exists():
            raise SystemExit(f"[ERROR] Part not found: {p}")
        append_docx(doc, p, page_break=True)

    apply_sau_fbe_page_setup(doc)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)

    print("[INFO] Written:", output)
    print("[INFO] Parts:", len(args.part))


if __name__ == "__main__":
    main()
